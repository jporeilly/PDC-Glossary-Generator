"""
suggester.py - core glossary-suggestion logic (importable).

Stages, all pure functions so both the CLI and the web app reuse them:
  harvest_ddl(path) / harvest_live(dsn)  -> {table: [column dicts]}
  suggest(tables)                        -> [row dicts]  (candidate terms)
  to_jsonl_records(rows, glossary_name)  -> [PDC objects] (glossary+cats+terms)

A "row" is the steward-facing review record (also what the UI edits).
"""
import os, re, json, uuid
import tagdict

DOMAIN = "General"
GEN_TS = "2026-06-18T12:00:00.000Z"

# ----------------------------------------------------------------- HARVEST
def harvest_ddl(path):
    """Parse a .sql DDL file into {table: [column dicts]} for offline scanning."""
    sql = open(path, encoding="utf-8").read()
    return harvest_ddl_text(sql)

def harvest_ddl_text(sql):
    """Parse DDL text (CREATE TABLE statements) into {table: [column dicts]},
       including foreign-key targets from inline `REFERENCES t(c)` and table-level
       `FOREIGN KEY (c) REFERENCES t(c)` so relationship edges can be drawn."""
    tables = {}
    ref_inline = re.compile(r"REFERENCES\s+(\w+)\s*\(\s*(\w+)\s*\)", re.I)
    fk_tablelevel = re.compile(
        r"FOREIGN\s+KEY\s*\(\s*(\w+)\s*\)\s*REFERENCES\s+(\w+)\s*\(\s*(\w+)\s*\)", re.I)
    for m in re.finditer(r"CREATE TABLE\s+(\w+)\s*\((.*?)\n\)\s*;", sql, re.S | re.I):
        tname, body = m.group(1), m.group(2)
        cols = []
        tbl_fks = {}  # column -> (ref_table, ref_col) from table-level constraints
        for fm in fk_tablelevel.finditer(body):
            tbl_fks[fm.group(1)] = (fm.group(2), fm.group(3))
        for raw in body.split("\n"):
            line = raw.strip().rstrip(",")
            if not line:
                continue
            comment = ""
            if "--" in line:
                line, comment = line.split("--", 1)
                line, comment = line.strip().rstrip(","), comment.strip()
            parts = line.split()
            if len(parts) < 2:
                continue
            col = parts[0]
            if col.upper() in {"PRIMARY", "FOREIGN", "CONSTRAINT", "UNIQUE", "CHECK", "REFERENCES"}:
                continue
            up = line.upper()
            ref = ref_inline.search(line)
            ref_table = ref.group(1) if ref else None
            ref_col = ref.group(2) if ref else None
            cols.append({"table": tname, "column": col, "type": parts[1],
                         "pk": "PRIMARY KEY" in up, "fk": "REFERENCES" in up,
                         "ref_table": ref_table, "ref_col": ref_col,
                         "notnull": "NOT NULL" in up, "unique": "UNIQUE" in up,
                         "comment": comment})
        # apply table-level FK constraints to their columns
        for c in cols:
            if c["column"] in tbl_fks:
                c["fk"] = True
                c["ref_table"], c["ref_col"] = tbl_fks[c["column"]]
        if cols:
            tables[tname] = cols
    return tables

def _harvest_oracle(conn, owner):
    """Oracle live scan. Oracle has no information_schema — the metadata lives in
    the ALL_* dictionary views — and python-oracledb uses :name binds, not %s.
    `owner` is the schema (Oracle schema == user, usually uppercase). Requires only
    SELECT on the dictionary views every account already has for its own objects."""
    tables = {}
    with conn.cursor() as cur:
        # columns (skip recycle-bin and system-generated $ objects)
        cur.execute(
            """SELECT table_name, column_name, data_type, column_id, nullable
               FROM all_tab_columns
               WHERE owner = :o
                 AND table_name NOT LIKE 'BIN$%'
                 AND table_name NOT LIKE '%$%'
               ORDER BY table_name, column_id""", o=owner)
        colrows = cur.fetchall()
        # primary keys
        cur.execute(
            """SELECT acc.table_name, acc.column_name
               FROM all_constraints ac
               JOIN all_cons_columns acc
                 ON acc.owner = ac.owner AND acc.constraint_name = ac.constraint_name
               WHERE ac.constraint_type = 'P' AND ac.owner = :o""", o=owner)
        pks = {(t, c) for t, c in cur.fetchall()}
        # foreign keys + their targets, position-aligned (handles composite keys)
        cur.execute(
            """SELECT a.table_name, a.column_name, pk.table_name, b.column_name
               FROM all_constraints c
               JOIN all_cons_columns a
                 ON a.owner = c.owner AND a.constraint_name = c.constraint_name
               JOIN all_constraints pk
                 ON pk.owner = c.r_owner AND pk.constraint_name = c.r_constraint_name
               JOIN all_cons_columns b
                 ON b.owner = pk.owner AND b.constraint_name = pk.constraint_name
                AND b.position = a.position
               WHERE c.constraint_type = 'R' AND c.owner = :o""", o=owner)
        fks, fkref = set(), {}
        for t, c, rt, rc in cur.fetchall():
            fks.add((t, c))
            fkref[(t, c)] = (rt, rc)
        # column comments
        cur.execute(
            """SELECT table_name, column_name, comments
               FROM all_col_comments
               WHERE owner = :o AND comments IS NOT NULL""", o=owner)
        comments = {(t, col): desc for t, col, desc in cur.fetchall()}
    for t, col, dt, pos, nullable in colrows:
        ref = fkref.get((t, col))
        tables.setdefault(t, []).append(
            {"table": t, "column": col, "type": dt,
             "pk": (t, col) in pks, "fk": (t, col) in fks,
             "ref_table": ref[0] if ref else None,
             "ref_col": ref[1] if ref else None,
             "notnull": (nullable == "N"), "unique": False,
             "comment": comments.get((t, col), "") or ""})
    return tables

def harvest_live(cfg, schema=None):
    """Live scan via a Python DB-API driver (see dbconn.py). cfg is a dict:
       {engine, host, port, database, schema, user, password, ssl}.
       Reads columns + keys + comments from information_schema (pg/mysql/mssql)
       or the ALL_* dictionary views (oracle — schema/owner defaults to the
       connecting user, uppercased)."""
    import dbconn
    eng = cfg.get("engine", "postgresql")
    schema = schema or cfg.get("schema") or ("public" if eng == "postgresql" else None)
    conn = dbconn._connect(cfg)
    tables = {}
    try:
        if eng == "oracle":
            owner = (schema or cfg.get("user") or "").strip().upper()
            if not owner:
                raise ValueError("Oracle scan needs a schema (owner) or a user to derive it from")
            return _harvest_oracle(conn, owner)
        with conn.cursor() as cur:
            # columns
            cur.execute(
                """SELECT table_name, column_name, data_type, ordinal_position,
                          (is_nullable='NO')
                   FROM information_schema.columns
                   WHERE table_schema = %s
                   ORDER BY table_name, ordinal_position""",
                (schema,))
            colrows = cur.fetchall()
            pks = fks = fkref = None
            if eng == "postgresql":
                # PRIMARY: read keys from pg_catalog, NOT information_schema. The
                # information_schema key views (key_column_usage / table_constraints)
                # are privilege-filtered and frequently come back EMPTY for a
                # connection user that can read columns but doesn't own the tables —
                # which is exactly why the diagram showed 0 PK / 0 FK. pg_catalog is
                # authoritative and reflects the constraints regardless of ownership.
                try:
                    cur.execute(
                        """SELECT c.relname, a.attname
                           FROM pg_index i
                           JOIN pg_class c ON c.oid = i.indrelid
                           JOIN pg_namespace n ON n.oid = c.relnamespace
                           JOIN pg_attribute a ON a.attrelid = c.oid AND a.attnum = ANY(i.indkey)
                           WHERE i.indisprimary AND n.nspname = %s""",
                        (schema,))
                    pks = {(t, c) for t, c in cur.fetchall()}
                    # FKs + their targets, column-aligned (handles composite keys),
                    # using generate_subscripts (older/portable) rather than multi-unnest.
                    cur.execute(
                        """SELECT c.relname, a.attname, cf.relname, af.attname
                           FROM pg_constraint con
                           JOIN pg_class c  ON c.oid  = con.conrelid
                           JOIN pg_namespace n ON n.oid = c.relnamespace
                           JOIN pg_class cf ON cf.oid = con.confrelid
                           JOIN generate_subscripts(con.conkey, 1) AS gs(i) ON true
                           JOIN pg_attribute a  ON a.attrelid  = con.conrelid  AND a.attnum  = con.conkey[gs.i]
                           JOIN pg_attribute af ON af.attrelid = con.confrelid AND af.attnum = con.confkey[gs.i]
                           WHERE con.contype = 'f' AND n.nspname = %s""",
                        (schema,))
                    fks, fkref = set(), {}
                    for t, c, rt, rc in cur.fetchall():
                        fks.add((t, c))
                        fkref[(t, c)] = (rt, rc)
                except Exception:
                    try:
                        conn.rollback()  # clear any aborted tx so the fallback can run
                    except Exception:
                        pass
                    pks = fks = fkref = None  # fall back to information_schema below
            if pks is None:
                # information_schema fallback (other engines, or if pg_catalog failed)
                cur.execute(
                    """SELECT kcu.table_name, kcu.column_name
                       FROM information_schema.table_constraints tc
                       JOIN information_schema.key_column_usage kcu
                         ON kcu.constraint_name = tc.constraint_name
                        AND kcu.table_schema = tc.table_schema
                       WHERE tc.constraint_type='PRIMARY KEY' AND tc.table_schema=%s""",
                    (schema,))
                pks = {(t, c) for t, c in cur.fetchall()}
                cur.execute(
                    """SELECT kcu.table_name, kcu.column_name
                       FROM information_schema.table_constraints tc
                       JOIN information_schema.key_column_usage kcu
                         ON kcu.constraint_name = tc.constraint_name
                        AND kcu.table_schema = tc.table_schema
                       WHERE tc.constraint_type='FOREIGN KEY' AND tc.table_schema=%s""",
                    (schema,))
                fks = {(t, c) for t, c in cur.fetchall()}
                fkref = {}
                try:
                    cur.execute(
                        """SELECT kcu.table_name, kcu.column_name,
                                  ccu.table_name AS ref_table, ccu.column_name AS ref_col
                           FROM information_schema.table_constraints tc
                           JOIN information_schema.key_column_usage kcu
                             ON kcu.constraint_name = tc.constraint_name
                            AND kcu.table_schema = tc.table_schema
                           JOIN information_schema.constraint_column_usage ccu
                             ON ccu.constraint_name = tc.constraint_name
                            AND ccu.table_schema = tc.table_schema
                           WHERE tc.constraint_type='FOREIGN KEY' AND tc.table_schema=%s""",
                        (schema,))
                    fkref = {(t, c): (rt, rc) for t, c, rt, rc in cur.fetchall()}
                except Exception:
                    fkref = {}
            # column comments (PostgreSQL)
            comments = {}
            if eng == "postgresql":
                cur.execute(
                    """SELECT c.relname, a.attname, d.description
                       FROM pg_description d
                       JOIN pg_class c ON c.oid = d.objoid
                       JOIN pg_attribute a ON a.attrelid = c.oid AND a.attnum = d.objsubid
                       JOIN pg_namespace n ON n.oid = c.relnamespace
                       WHERE n.nspname = %s AND d.objsubid > 0""",
                    (schema,))
                comments = {(t, col): desc for t, col, desc in cur.fetchall()}
        for t, col, dt, pos, notnull in colrows:
            ref = fkref.get((t, col))
            tables.setdefault(t, []).append(
                {"table": t, "column": col, "type": dt,
                 "pk": (t, col) in pks, "fk": (t, col) in fks,
                 "ref_table": ref[0] if ref else None,
                 "ref_col": ref[1] if ref else None,
                 "notnull": bool(notnull), "unique": False,
                 "comment": comments.get((t, col), "") or ""})
    finally:
        conn.close()
    return tables

def schema_graph(tables):
    """Shape scanned {table: [col dicts]} into an ER graph for the diagram:
       tables[{name, columns[{name,type,pk,fk,notnull,ref_table,ref_col}], pk_count,
       fk_count}] and relationships[{from,from_col,to,to_col}] (only FKs whose target
       table is present in the scan are marked resolved)."""
    names = set(tables.keys())
    out_tables, rels = [], []
    for tname, cols in tables.items():
        out_cols = []
        for c in cols:
            out_cols.append({
                "name": c.get("column"), "type": c.get("type", ""),
                "pk": bool(c.get("pk")), "fk": bool(c.get("fk")),
                "notnull": bool(c.get("notnull")),
                "ref_table": c.get("ref_table"), "ref_col": c.get("ref_col"),
                "comment": c.get("comment", "") or ""})
            if c.get("fk") and c.get("ref_table"):
                rels.append({"from": tname, "from_col": c.get("column"),
                             "to": c.get("ref_table"), "to_col": c.get("ref_col"),
                             "resolved": c.get("ref_table") in names})
        out_tables.append({
            "name": tname, "columns": out_cols,
            "pk_count": sum(1 for x in out_cols if x["pk"]),
            "fk_count": sum(1 for x in out_cols if x["fk"]),
            "col_count": len(out_cols)})
    out_tables.sort(key=lambda t: t["name"])
    return {"tables": out_tables, "relationships": rels,
            "table_count": len(out_tables),
            "rel_count": sum(1 for r in rels if r["resolved"])}


def keymap_from_tables(tables):
    """Reduce scanned {table: [col dicts]} to the keys we'd want set on the DB:
       {table: {pk:[cols], fks:[{col, ref_table, ref_col}]}}."""
    km = {}
    for tname, cols in tables.items():
        pk = [c["column"] for c in cols if c.get("pk")]
        fks = [{"col": c["column"], "ref_table": c.get("ref_table"), "ref_col": c.get("ref_col")}
               for c in cols if c.get("fk") and c.get("ref_table") and c.get("ref_col")]
        if pk or fks:
            km[tname] = {"pk": pk, "fks": fks}
    return km


def sample_distinct_values(cfg, sources, limit=200):
    """Live-data probe for the duplicate-group recommender: sample up to `limit`
       DISTINCT non-null values for each 'schema.table.column' source. Direct
       value overlap between two same-named columns is the strongest same-vs-
       different-concept evidence there is (better than cached profile shapes,
       because it compares the actual populations). Returns {source: [values]};
       sources that fail to read are simply absent. Postgres/MySQL/MSSQL via
       dbconn (Oracle uses FETCH FIRST)."""
    import dbconn
    eng = cfg.get("engine", "postgresql")
    conn = dbconn._connect(cfg)
    out = {}
    try:
        with conn.cursor() as cur:
            for src in sources or []:
                bits = str(src).strip().split(".")
                if len(bits) < 3:
                    continue
                schema, table, col = bits[-3], bits[-2], bits[-1]
                if not all(re.fullmatch(r"[A-Za-z0-9_$]+", x) for x in (schema, table, col)):
                    continue                      # identifiers only — never quote-inject
                n = max(1, min(int(limit), 1000))
                if eng == "oracle":
                    q = (f'SELECT DISTINCT "{col.upper()}" FROM "{schema.upper()}"."{table.upper()}" '
                         f'WHERE "{col.upper()}" IS NOT NULL FETCH FIRST {n} ROWS ONLY')
                else:
                    q = (f'SELECT DISTINCT "{col}" FROM "{schema}"."{table}" '
                         f'WHERE "{col}" IS NOT NULL LIMIT {n}')
                try:
                    cur.execute(q)
                    out[src] = [str(r[0]) for r in cur.fetchall()]
                except Exception:
                    try:
                        conn.rollback()
                    except Exception:
                        pass
    finally:
        try:
            conn.close()
        except Exception:
            pass
    return out


def apply_keys_live(cfg, schema, keymap, dry_run=True):
    """Add the PRIMARY KEY / FOREIGN KEY constraints in `keymap` to a live PostgreSQL
       schema via ALTER TABLE, so PDC's catalog ingest (and our own scan) can read
       them. Idempotent: existing keys are skipped. Each statement runs in its own
       sub-transaction so one failure (e.g. an orphan FK value) doesn't block the
       rest. Returns a per-statement report; dry_run just returns the planned SQL."""
    import dbconn
    eng = cfg.get("engine", "postgresql")
    schema = schema or cfg.get("schema") or "public"
    if eng != "postgresql":
        raise RuntimeError("Writing keys is currently supported for PostgreSQL only.")
    conn = dbconn._connect(cfg)
    stmts = []
    skipped_pk = skipped_fk = 0
    try:
        with conn.cursor() as cur:
            cur.execute(
                """SELECT c.relname FROM pg_constraint con
                   JOIN pg_class c ON c.oid = con.conrelid
                   JOIN pg_namespace n ON n.oid = c.relnamespace
                   WHERE con.contype='p' AND n.nspname=%s""", (schema,))
            haspk = {r[0] for r in cur.fetchall()}
            cur.execute(
                """SELECT c.relname, a.attname FROM pg_constraint con
                   JOIN pg_class c ON c.oid = con.conrelid
                   JOIN pg_namespace n ON n.oid = c.relnamespace
                   JOIN generate_subscripts(con.conkey, 1) AS gs(i) ON true
                   JOIN pg_attribute a ON a.attrelid = con.conrelid AND a.attnum = con.conkey[gs.i]
                   WHERE con.contype='f' AND n.nspname=%s""", (schema,))
            fkcols = {(t, c) for t, c in cur.fetchall()}
            cur.execute(
                "SELECT table_name FROM information_schema.tables WHERE table_schema=%s", (schema,))
            existing = {r[0] for r in cur.fetchall()}

            def q(t):
                return '"%s"."%s"' % (schema, t)

            plan = []
            for t, km in keymap.items():
                if t not in existing:
                    continue
                if km["pk"]:
                    if t in haspk:
                        skipped_pk += 1
                    else:
                        cols = ", ".join('"%s"' % c for c in km["pk"])
                        plan.append(("pk", t, 'ALTER TABLE %s ADD PRIMARY KEY (%s)' % (q(t), cols)))
                for fk in km["fks"]:
                    if fk["ref_table"] not in existing:
                        continue
                    if (t, fk["col"]) in fkcols:
                        skipped_fk += 1
                        continue
                    cn = "%s_%s_fkey" % (t, fk["col"])
                    plan.append(("fk", t, 'ALTER TABLE %s ADD CONSTRAINT "%s" FOREIGN KEY ("%s") REFERENCES %s ("%s")'
                                 % (q(t), cn, fk["col"], q(fk["ref_table"]), fk["ref_col"])))

            if dry_run:
                for kind, t, sql in plan:
                    stmts.append({"kind": kind, "table": t, "sql": sql, "status": "pending"})
            else:
                for kind, t, sql in plan:
                    try:
                        cur.execute(sql)
                        conn.commit()
                        stmts.append({"kind": kind, "table": t, "sql": sql, "status": "applied"})
                    except Exception as e:
                        conn.rollback()
                        stmts.append({"kind": kind, "table": t, "sql": sql, "status": "error",
                                      "message": str(e).splitlines()[0][:200]})
    finally:
        conn.close()
    return {"schema": schema, "statements": stmts, "dry_run": dry_run,
            "applied": sum(1 for s in stmts if s["status"] == "applied"),
            "errors": sum(1 for s in stmts if s["status"] == "error"),
            "pending": sum(1 for s in stmts if s["status"] == "pending"),
            "skipped_pk": skipped_pk, "skipped_fk": skipped_fk,
            "pk_planned": sum(1 for s in stmts if s["kind"] == "pk"),
            "fk_planned": sum(1 for s in stmts if s["kind"] == "fk")}

# ---- value-level data profiling: sample real data to determine sensitivity/CDE ----
RX_EMAIL = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
RX_SSN   = re.compile(r"^\d{3}-\d{2}-\d{4}$")
RX_CC    = re.compile(r"^\d{13,19}$")
RX_PHONE = re.compile(r"^[+(]?[\d][\d\s().-]{6,}$")
RX_ZIP   = re.compile(r"^\d{5}(-\d{4})?$")
RX_DATE  = re.compile(r"^\d{4}-\d{2}-\d{2}|^\d{1,2}/\d{1,2}/\d{2,4}")
RX_DEC   = re.compile(r"^-?\d+\.\d+$")

def _value_signature(s):
    """PDC-style position signature of one value: digits->n, upper->A, lower->a,
       common separators kept literally. None for long/exotic values."""
    if not s or len(s) > 32:
        return None
    out = []
    for ch in s:
        if ch.isdigit(): c = "n"
        elif ch.isalpha(): c = "A" if ch.isupper() else "a"
        elif ch in "-_./ :#": c = ch
        else: return None
        out.append(c)
    return "".join(out)

def _induce_pattern(strs):
    r"""Learn a value format from sampled data: when >=90% of values share one
       position signature (e.g. AAA-nnnnn for CPC-84120), derive an anchored
       regex — a stable literal prefix is kept verbatim (^CPC-\d{5}$), the rest
       generalizes by character class. This is the evidence a Data Pattern
       method needs, so it flows into the row and the Registry `detect` list.
       Returns (signature, regex, share) or (None, None, 0)."""
    sigs = {}
    for v in strs:
        g = _value_signature(v)
        if g:
            sigs.setdefault(g, []).append(v)
    if not sigs:
        return None, None, 0
    sig, vals = max(sigs.items(), key=lambda kv: len(kv[1]))
    share = len(vals) / len(strs)
    if share < 0.9 or len(vals) < 5 or len(sig) < 4:
        return None, None, 0
    if "n" not in sig or len(set(sig)) < 2:
        return None, None, 0        # want structured codes, not plain words/numbers
    prefix = os.path.commonprefix(vals)
    while prefix and (prefix[-1].isdigit() or _value_signature(prefix) is None):
        prefix = prefix[:-1]        # never let variance digits leak into the literal
    if len(prefix) < 2:
        prefix = ""
    rest = sig[len(prefix):]
    parts, i = [], 0
    while i < len(rest):
        j = i
        while j < len(rest) and rest[j] == rest[i]:
            j += 1
        k, c = j - i, rest[i]
        if c == "n":   parts.append(r"\d{%d}" % k if k > 1 else r"\d")
        elif c == "A": parts.append("[A-Z]{%d}" % k if k > 1 else "[A-Z]")
        elif c == "a": parts.append("[a-z]{%d}" % k if k > 1 else "[a-z]")
        else:          parts.append(re.escape(c) * k)
        i = j
    rx = "^" + re.escape(prefix) + "".join(parts) + "$"
    try:
        crx = re.compile(rx)
    except re.error:
        return None, None, 0
    ok = sum(1 for v in strs if crx.match(v)) / len(strs)
    if ok < 0.9:
        return None, None, 0
    return sig, rx, ok

def _profile_values(name, vals, sample_n):
    """Infer pii/sensitivity/uniqueness/type from a column's sampled values.
       Also returns DQ signals: completeness (non-empty/sampled) and, where a
       pattern/type is detected, validity (share of values conforming)."""
    strs = [str(v).strip() for v in vals if v is not None and str(v).strip() != ""]
    completeness = round(len(strs) / sample_n, 3) if sample_n else 0
    if not strs:
        return {"uniq": 0, "completeness": completeness, "kind": "empty", "reason": "Profiled: all null/blank"}
    n = len(strs)
    def frac(rx):
        """Fraction of the sampled values that match regex `rx`."""
        return sum(1 for s in strs if rx.match(s)) / n
    distinct = len(set(strs)); uniq = distinct / n
    avg_digits = sum(sum(ch.isdigit() for ch in s) for s in strs) / n
    has_sep = sum(1 for s in strs if any(c in s for c in "-() ")) / n
    base = {"uniq": uniq, "completeness": completeness}
    if frac(RX_EMAIL) >= 0.6:
        return {**base, "pii": "CONTACT_INFO", "sensitivity": "HIGH", "confidence": "High",
                "reason": "Profiled: email values", "kind": "email", "valid": round(frac(RX_EMAIL), 3)}
    if frac(RX_SSN) >= 0.6:
        return {**base, "pii": "PERSONAL_NAME", "sensitivity": "HIGH", "confidence": "High",
                "reason": "Profiled: SSN-format values", "kind": "ssn", "valid": round(frac(RX_SSN), 3)}
    if frac(RX_CC) >= 0.6 and avg_digits >= 13:
        return {**base, "pii": "FINANCIAL", "sensitivity": "HIGH", "confidence": "High",
                "reason": "Profiled: card/account-number values", "kind": "card", "valid": round(frac(RX_CC), 3)}
    if frac(RX_DATE) >= 0.6:
        return {**base, "confidence": "Medium", "reason": "Profiled: date values", "kind": "date", "valid": round(frac(RX_DATE), 3)}
    if frac(RX_ZIP) >= 0.7:
        return {**base, "pii": "ADDRESS_INFO", "sensitivity": "MEDIUM", "confidence": "High",
                "reason": "Profiled: postal-code values", "kind": "zip", "valid": round(frac(RX_ZIP), 3)}
    if frac(RX_PHONE) >= 0.6 and 7 <= avg_digits <= 15 and has_sep >= 0.3:
        return {**base, "pii": "CONTACT_INFO", "sensitivity": "MEDIUM", "confidence": "High",
                "reason": "Profiled: phone-format values", "kind": "phone", "valid": round(frac(RX_PHONE), 3)}
    if distinct <= 12 and n >= 10:
        return {**base, "confidence": "Medium", "kind": "enum",
                "reason": f"Profiled: low cardinality ({distinct} distinct - reference-data candidate)",
                "enum": sorted(set(strs))[:12]}
    sig, rx, share = _induce_pattern(strs)
    if sig:
        return {**base, "confidence": "High",
                "kind": "identifier" if uniq >= 0.95 else "code",
                "signature": sig, "pattern": rx, "valid": round(share, 3),
                "reason": f"Profiled: {int(share * 100)}% of values share position signature {sig}"}
    if uniq >= 0.95 and n >= 5 and frac(RX_DEC) < 0.5:
        return {**base, "confidence": "High", "reason": "Profiled: near-unique values (likely identifier)",
                "kind": "identifier"}
    dec = frac(RX_DEC)
    if dec >= 0.5:
        return {**base, "reason": "Profiled", "kind": "decimal", "valid": round(dec, 3)}
    return {**base, "reason": "Profiled", "kind": "value"}

def profile_live(cfg, tables, schema=None, sample_size=80):
    """Sample rows per table and attach a `profile` dict to each column. Best-effort;
       columns/tables that can't be sampled are left name-based."""
    import dbconn
    eng = cfg.get("engine", "postgresql")
    schema = schema or cfg.get("schema") or "public"
    conn = dbconn._connect(cfg)
    try:
        with conn.cursor() as cur:
            for tname, cols in tables.items():
                try:
                    if eng == "sqlserver":
                        cur.execute(f'SELECT TOP {sample_size} * FROM "{schema}"."{tname}"')
                    else:
                        cur.execute(f'SELECT * FROM "{schema}"."{tname}" LIMIT {sample_size}')
                    names = [d[0] for d in cur.description]
                    rows = cur.fetchall()
                except Exception:
                    conn.rollback() if hasattr(conn, "rollback") else None
                    continue
                for col in cols:
                    if col["column"] in names:
                        idx = names.index(col["column"])
                        vals = [r[idx] for r in rows]
                        col["profile"] = _profile_values(col["column"], vals, len(rows))
                        seen, ex = set(), []
                        for v in vals:
                            if v is None: continue
                            s = str(v).strip()
                            if s and s not in seen:
                                seen.add(s); ex.append(s[:40])
                            if len(ex) >= 3: break
                        col["examples"] = ex
    finally:
        conn.close()
    return tables

def discover(cfg, schema=None, sample_size=100):
    """Full data-discovery profile per table/column: row counts, completeness,
       distinct/uniqueness, sensitivity/PII/CDE, detected kind and example values.
       Mirrors the dimensions PDC's column profiling shows, for side-by-side comparison."""
    import dbconn
    schema = schema or cfg.get("schema") or "public"
    tables = harvest_live(cfg, schema)
    try:
        profile_live(cfg, tables, schema, sample_size)
    except Exception:
        pass
    rows = suggest(tables, schema)
    bykey = {}
    for r in rows:
        sc = r["Source_Column"].split(";")[0].strip().split(".")
        if len(sc) >= 3:
            bykey[(sc[-2], sc[-1])] = r
    conn = dbconn._connect(cfg)
    out, summary = [], {"tables": 0, "columns": 0, "rows": 0, "pii": 0, "cde": 0,
                        "empty": 0, "sensitivity": {"HIGH": 0, "MEDIUM": 0, "LOW": 0},
                        "pk_cols": 0, "fk_cols": 0, "classified": 0, "db_bytes": 0,
                        "avg_completeness": 0}
    comp_sum = comp_n = 0
    try:
        with conn.cursor() as cur:
            try:
                cur.execute("SELECT pg_database_size(current_database())")
                summary["db_bytes"] = cur.fetchone()[0] or 0
            except Exception:
                try: conn.rollback()
                except Exception: pass
            for tname, cols in tables.items():
                sel = ["COUNT(*)"]
                for c in cols:
                    q = '"' + c["column"].replace('"', '') + '"'
                    sel += [f"COUNT({q})", f"COUNT(DISTINCT {q})"]
                try:
                    cur.execute(f'SELECT {", ".join(sel)} FROM "{schema}"."{tname}"')
                    agg = cur.fetchone()
                except Exception:
                    try: conn.rollback()
                    except Exception: pass
                    agg = None
                tbytes = 0
                try:
                    cur.execute("SELECT pg_total_relation_size(%s::regclass)", (f'"{schema}"."{tname}"',))
                    tbytes = cur.fetchone()[0] or 0
                except Exception:
                    try: conn.rollback()
                    except Exception: pass
                total = (agg[0] if agg else 0) or 0
                colout = []
                for i, c in enumerate(cols):
                    nn = (agg[1 + i * 2] if agg else 0) or 0
                    dd = (agg[2 + i * 2] if agg else 0) or 0
                    sr = bykey.get((tname, c["column"]), {})
                    prof = c.get("profile") or {}
                    sens = sr.get("Sensitivity", "LOW"); pii = sr.get("PII_Category", ""); cde = sr.get("Critical_Data_Element", "No")
                    completeness = round(nn / total, 3) if total else 0
                    colout.append({"column": c["column"], "type": c["type"], "pk": c["pk"], "fk": c["fk"],
                                   "non_null": nn, "distinct": dd,
                                   "completeness": completeness,
                                   "uniqueness": round(dd / nn, 3) if nn else 0,
                                   "sensitivity": sens, "pii": pii, "cde": cde,
                                   "kind": prof.get("kind", ""), "examples": c.get("examples", []),
                                   "term": sr.get("Term", ""), "confidence": sr.get("Confidence", "")})
                    summary["columns"] += 1
                    if pii: summary["pii"] += 1
                    if cde == "Yes": summary["cde"] += 1
                    if c["pk"]: summary["pk_cols"] += 1
                    if c["fk"]: summary["fk_cols"] += 1
                    if pii or sens != "LOW": summary["classified"] += 1
                    if total:
                        comp_sum += completeness; comp_n += 1
                    if sens in summary["sensitivity"]: summary["sensitivity"][sens] += 1
                out.append({"name": tname, "rows": total, "bytes": tbytes,
                            "empty": total == 0, "columns": colout})
                summary["tables"] += 1; summary["rows"] += total
                if total == 0: summary["empty"] += 1
    finally:
        conn.close()
    summary["avg_completeness"] = round(comp_sum / comp_n, 3) if comp_n else 0
    summary["largest_tables"] = sorted(
        [{"name": t["name"], "rows": t["rows"], "bytes": t["bytes"]} for t in out],
        key=lambda t: t["rows"], reverse=True)[:5]
    return {"schema": schema, "tables": out, "summary": summary}

# ----------------------------------------------------------------- SUGGEST
def _load_domain_pack():
    """Optionally load scenario vocabulary so the engine stays generic. Looks at
       $GLOSSARY_DOMAIN_PACK, else domain_pack.json beside this module. Returns {}
       when absent. Recognised keys: table_category, table_terms, cat_keywords,
       abbreviations, category_definitions. See domain_packs/*.example.json."""
    path = os.environ.get("GLOSSARY_DOMAIN_PACK") or \
        os.path.join(os.path.dirname(__file__), "domain_pack.json")
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

_PACK = _load_domain_pack()

# Physical table -> business category. Empty by default (table names are
# scenario-specific); populate via a domain pack for your schema.
TABLE_CATEGORY = dict(_PACK.get("table_category", {}))

# Table-level glossary terms — the ONE term per physical table that feeds the
# table's Trust Score "glossary term assigned" input. Columns carry their own
# terms (linked to the column); the TABLE needs a term on the table entity itself,
# because PDC reads the table's own businessTerms[] for that input. Curated names
# come from the domain pack (if any); otherwise a name is derived per table.
TABLE_TERMS = dict(_PACK.get("table_terms", {}))

def _singularize(word):
    w = (word or "")
    if w.endswith("ies") and len(w) > 3:
        return w[:-3] + "y"
    if w.endswith("ses") and len(w) > 3:
        return w[:-2]
    if w.endswith("s") and not w.endswith("ss"):
        return w[:-1]
    return w

def table_term_name(table):
    """The table-level term for a physical table — the term linked to the TABLE
       entity so its Trust Score gets the 'term assigned' input. Uses the domain pack's
       names where known, else derives '<Singular Table> Record'."""
    t = (table or "").strip().lower()
    if t in TABLE_TERMS:
        return TABLE_TERMS[t]
    human = humanize(_singularize(t))
    return f"{human} Record" if human else "Record"

def table_term_rows(tables, col_rows=None):
    """One table-level "record" term per scanned table, as a CONCEPTUAL glossary-only
       row: empty Source_Column means it is created in the glossary but never auto-linked.
       The Data Steward links each term to its table by hand to feed that table's Trust
       Score "term assigned" input — kept a manual task until Pentaho clarifies the Data
       Quality direction. Sensitivity inherits the table's highest column sensitivity so a
       record term is at least as sensitive as the data it represents."""
    col_rows = col_rows or []
    tmax = {}
    for r in col_rows:
        src = str(r.get("Source_Column", "")).split(";")[0].strip()
        parts = src.split(".")
        if len(parts) >= 3:
            tmax[parts[1]] = max(tmax.get(parts[1], 0), SENS_RANK.get(r.get("Sensitivity", "LOW"), 0))
    rows = []
    for tname in tables:
        term = table_term_name(tname)
        human_tbl = humanize(_singularize(tname)).lower()
        rows.append({
            "Keep": "Y", "Category": categorize(tname), "Term": term, "Source_Column": "",
            "Source_Table": tname,
            "Definition": f"A single {human_tbl} record — the table-level business term for the {tname} table.",
            "Purpose": f"Linked to the {tname} table at Apply (table roll-up) to give its Trust Score the assigned-term input.",
            "Sensitivity": RANK_SENS[tmax.get(tname, SENS_RANK.get("LOW", 0))],
            "PII_Category": "", "Critical_Data_Element": "No",
            "Abbreviation": _abbrev(term), "Suggested_Tags": ";".join(
                ["record", "table-level"] + tagdict.category_tags().get(categorize(tname), [])),
            "Suggested_Rating": 0, "Source_Ratings": {},
            "Suggested_Quality": None, "Source_Quality_Dims": {},
            "Status": "Draft", "Confidence": "",
            "Suggested_Reason": "Table-level term — Steward links it to the table by hand to feed the Trust Score.",
            "LLM_Enriched": "No", "Map": "No"})
    return rows

# keyword fallback for tables not in the explicit map (order matters - first hit wins)
CAT_KEYWORDS = [
    ("billing", "Billing & Rates"), ("rate", "Billing & Rates"), ("invoice", "Billing & Rates"),
    ("usage", "Usage"), ("consumption", "Usage"), ("meter", "Usage"),
    ("alert", "Governance"), ("audit", "Governance"), ("compliance", "Governance"), ("policy", "Governance"),
    ("document", "Records & Documents"), ("record", "Records & Documents"),
    ("customer", "Customer"), ("account", "Customer"),
] + [tuple(x) for x in _PACK.get("cat_keywords", [])]

def categorize(tname):
    """Map a physical table name to a business glossary category."""
    if tname in TABLE_CATEGORY:
        return TABLE_CATEGORY[tname]
    t = tname.lower()
    for kw, cat in CAT_KEYWORDS:
        if kw in t:
            return cat
    return "Uncategorized"

# ---------------------------------------------------------------------------
# PII / sensitivity classification by COLUMN NAME.
#
# Each rule is a tuple:  (match, exclude, pii_category, sensitivity, tags)
#   match        regex tested against the lower-cased column name
#   exclude      regex that VETOES the match (e.g. "name" matches, but not
#                "system_name"/"file_name" which aren't personal names)
#   pii_category PDC PII bucket the column maps to (FINANCIAL, CONTACT_INFO…)
#   sensitivity  HIGH / MEDIUM / LOW assigned when this rule wins
#   tags         seed tags merged into the term's tag set
#
# Order matters: the FIRST rule that matches wins, so the most specific /
# highest-risk patterns are listed first (account number, SSN, email…) and the
# broad/low-risk ones last. A column that matches nothing is LOW with no PII.
# Value-level profiling (when "Sample values" is on) can OVERRIDE this name-based
# guess with what's actually in the data — see classify_values() above.
# ---------------------------------------------------------------------------
PII_RULES = [
    (r"account_number|acct", None, "FINANCIAL", "HIGH", ["pii", "financial"]),
    (r"ssn|social_security", None, "GOVERNMENT_ID", "HIGH", ["pii"]),
    (r"tax_?id|\bein\b|passport|driver_?licen[cs]e", None, "GOVERNMENT_ID", "HIGH", ["pii"]),
    (r"email|e_mail", None, "CONTACT_INFO", "HIGH", ["pii"]),
    (r"birth|dob|date_of_birth", None, "DEMOGRAPHIC", "HIGH", ["pii"]),
    (r"phone|mobile|telephone", None, "CONTACT_INFO", "MEDIUM", ["pii"]),
    (r"name", r"system|report|file|plan|type|source", "PERSONAL_NAME", "MEDIUM", ["pii"]),
    (r"address|street", None, "ADDRESS_INFO", "HIGH", ["pii"]),
    (r"(?<![a-z])(city|county|zip|postal|province|state)(?![a-z])", None, "ADDRESS_INFO", "MEDIUM", []),
    (r"amount|charge|tax|due|paid|balance", None, "FINANCIAL", "LOW", ["financial"]),
]
ABBREV = {"number": "No.", "identifier": "ID", "amount": "Amt", "account": "Acct",
          "address": "Addr", "quantity": "Qty", "percentage": "Pct"}
SKIP = re.compile(r"^(last_updated|created_date|created_at|updated_at)$", re.I)

# Token-level expansion for cryptic/abbreviated column names, so a horrible name
# like "cust_acct_no" becomes "Customer Account Number" rather than "Cust Acct No".
# Keys are matched per underscore-separated token (case-insensitive, exact token).
# Values are stored already-cased and inserted verbatim (so "ID"/"SSN" stay upper and
# multi-word expansions like "Date of Birth" keep their small words lower-case).
# Conservative on purpose: only well-known abbreviations a steward would expand by
# hand. Anything not listed falls through to plain Title-Casing, and every result is
# still only a *suggestion* the reviewer can edit.
EXPAND = {
    # identity / generic
    "id": "ID", "no": "Number", "num": "Number", "nbr": "Number", "cd": "Code",
    "nm": "Name", "desc": "Description", "ref": "Reference", "seq": "Sequence",
    "flg": "Flag", "ind": "Indicator", "stat": "Status", "sts": "Status",
    "qty": "Quantity", "amt": "Amount", "pct": "Percent", "avg": "Average",
    "tot": "Total", "bal": "Balance", "min": "Minimum", "max": "Maximum",
    # people / contact
    "cust": "Customer", "acct": "Account", "addr": "Address", "fname": "First Name",
    "lname": "Last Name", "dob": "Date of Birth", "ssn": "SSN", "tel": "Telephone",
    "ph": "Phone", "phn": "Phone", "email": "Email", "zip": "ZIP",
    # time
    "dt": "Date", "ts": "Timestamp", "yr": "Year", "mo": "Month", "qtr": "Quarter",
    "wk": "Week", "hr": "Hour",
    # finance / billing
    "txn": "Transaction", "trans": "Transaction", "inv": "Invoice", "pmt": "Payment",
    "freq": "Frequency", "curr": "Currency",
    # location / general
    "svc": "Service", "sys": "System", "loc": "Location", "geo": "Geographic",
    "lat": "Latitude", "lon": "Longitude", "lng": "Longitude",
}
EXPAND.update(_PACK.get("abbreviations", {}))

def humanize(col):
    """Turn a snake_case identifier into a human-readable Title Case label, expanding
       well-known abbreviations (see EXPAND) so cryptic column names still read well."""
    s = re.sub(r"\s+", " ", re.sub(r"[_]+", " ", col).strip())
    out = []
    for w in s.split():
        rep = EXPAND.get(w.lower())
        if rep is not None:
            out.append(rep)                 # already-cased expansion, inserted verbatim
        else:
            out.append(w if w.isupper() else w.capitalize())
    return " ".join(out)

def _abbrev(name):
    """Derive a short uppercase abbreviation from a term name."""
    for w in name.lower().split():
        if w in ABBREV:
            return ABBREV[w]
    return ""

def classify(col):
    """Classify a column name into (pii_category, sensitivity, tags)."""
    cl = col.lower()
    for pat, excl, cat, sens, tags in PII_RULES:
        if re.search(pat, cl) and not (excl and re.search(excl, cl)):
            return cat, sens, list(tags)
    return None, "LOW", []

def define(c):
    """Compose a plain-language DEFINITION for a column.

    Priority (best evidence first):
      1. The database COMMENT on the column, if the DBA wrote one — this is the
         authoritative business meaning, so it's used verbatim. A short
         comma-list comment is treated as an enumeration ("Valid values: …").
      2. A primary key  -> "Unique identifier for a <entity> record."
      3. A foreign key  -> "Reference linking this record to its related <ref>."
      4. Fallback       -> a neutral template from the humanised name.
    The LLM-enrich step can later rewrite any of these into richer prose; this
    function only guarantees every term ships with a sensible definition.
    """
    human_tbl = humanize(c["table"]).rstrip("s")   # "customers" -> "Customer"
    name = humanize(c["column"])                    # "service_address" -> "Service Address"
    if c["comment"]:
        # a comma-separated comment under ~90 chars reads as an enum of valid values
        if "," in c["comment"] and len(c["comment"].split(",")) >= 2 and len(c["comment"]) < 90:
            return f"{name} for a {human_tbl.lower()} record. Valid values: {c['comment']}."
        return c["comment"].rstrip(".") + "."       # use the DBA's comment as-is
    if c["pk"]:
        return f"Unique identifier for a {human_tbl.lower()} record."
    if c["fk"]:
        ref = humanize(c["column"]).replace(" ID", "").replace(" Id", "").strip()
        return f"Reference linking this record to its related {ref.lower()}."
    return f"{name} associated with a {human_tbl.lower()} record."

def purpose(c, category, name, pii):
    """A business 'why this matters / how it's used' sentence (the Purpose field)."""
    if c["pk"]:
        ent = humanize(c["table"]).rstrip("s").lower()
        return f"Uniquely identifies each {ent} for joins, lineage, and record integrity."
    if c["fk"]:
        ref = name.replace(" ID", "").replace(" Id", "").strip().lower() or "related record"
        return f"Links records to their related {ref} for analysis and lineage."
    if pii in ("PERSONAL_NAME", "CONTACT_INFO"):
        return "Identifies and contacts the customer; governed for privacy and regulatory compliance."
    if pii == "ADDRESS_INFO":
        return "Locates the customer for service and correspondence; governed for privacy."
    if pii == "FINANCIAL":
        return "Supports billing, revenue reporting, and financial reconciliation."
    return {**{
        "Customer": "Maintains customer records for service, billing, and communication.",
        "Billing & Rates": "Supports billing, rate calculation, and revenue reporting.",
        "Usage": "Measures usage or consumption for billing, forecasting, and analysis.",
        "Governance": "Supports governance, alerting, and compliance tracking.",
        "Records & Documents": "Stores supporting documents for reference and compliance.",
    }, **_PACK.get("category_definitions", {})}.get(
        category, f"Provides {category.lower()} context for reporting, governance, and discovery.")

def _slug(s):
    """Return a lower-cased, id-safe slug of a string."""
    return re.sub(r"[^a-z0-9]+", "-", s.lower()).strip("-")

# Critical Data Element: governed with the highest care. Inferred from keys, high
# sensitivity, financial/identity PII, and critical business/compliance/safety terms.
# Always reviewable by the Data Steward in the grid.
CDE_PATTERNS = re.compile(
    r"(account.?number|\bssn\b|social.?security|tax.?id|\bein\b|"
    r"licen[cs]e|permit|meter.?(id|number|no)|balance|amount.?(due|owed)|"
    r"complian|violation|lead.?(level|ppb)?|contaminant|ph.?level|turbidity)", re.I)

def is_cde(name, sens, pii, key_like):
    """Heuristic test for a Critical Data Element — a field governed with the
    highest care (it materially affects billing, identity, safety or compliance).
    A column is flagged CDE when ANY of these hold; the steward can always override
    the call in the grid:
      - it's a primary key (an identity/anchor column),
      - it's HIGH sensitivity (already classified as risky),
      - it's financial or personal-name PII,
      - or its name matches a critical pattern (account number, meter id, balance,
        compliance/violation, lead level, pH, turbidity… — see CDE_PATTERNS).
    """
    if key_like:                       # primary key = identity anchor (plain FKs excluded)
        return True
    if sens == "HIGH":                 # already deemed high-risk by classification
        return True
    if pii in ("FINANCIAL", "PERSONAL_NAME"):   # money + direct identifiers
        return True
    return bool(CDE_PATTERNS.search(name or ""))  # domain-critical name patterns

# --------------------------------------------------------------------------- #
#  Meaningful, controlled tags — sourced from the per-company TAG DICTIONARY
#  (tagdict.py): a persisted allow-list + name->tag rules, seeded from the domain
#  and grown from scans. suggest_tags() reads the *live* dictionary, so tags stay
#  consistent with what the Registry (and, downstream, the Policy Generator) uses.
# --------------------------------------------------------------------------- #


def suggest_tags(category, sens, pii, cde, is_key, base_tags=None, name="", term=""):
    """Build a deterministic, meaningful, de-duplicated tag set for a term, drawn
    from the controlled tag dictionary (allow-list + rules):
      - PII type          -> privacy / contact / location / financial tags,
      - name/term/category-> domain tags via the dictionary's rules,
      - a meaningful category tag (dictionary category_tags), not just the slug,
      - HIGH sens -> 'maskable', CDE -> 'cde', key -> 'identifier'.
    Everything stays within the dictionary's vocabulary so tags can't drift.
    """
    t = list(base_tags or [])
    if pii == "PERSONAL_NAME":  t += ["pii", "personal-data", "direct-identifier", "privacy"]
    elif pii == "CONTACT_INFO": t += ["pii", "contact", "privacy"]
    elif pii == "ADDRESS_INFO": t += ["pii", "location", "privacy"]
    elif pii == "FINANCIAL":    t += ["financial", "sensitive"]

    hay = " ".join([str(name or ""), str(term or ""), str(category or "")])
    for rx, tags in tagdict.compiled_rules():
        if rx.search(hay):
            t += tags

    cat_tags = tagdict.category_tags().get(category)
    t += cat_tags if cat_tags else ([_slug(category)] if category else [])

    if sens == "HIGH":                              t.append("maskable")
    if str(cde).lower() == "yes" or cde is True:    t.append("cde")
    if is_key:                                      t.append("identifier")

    seen, out = set(), []          # standardised lower-case, de-duped, order kept
    for x in t:
        k = str(x or "").strip().lower()
        if k and k not in seen:
            seen.add(k); out.append(k)
    return out[:7]


def retag_rows(rows):
    """Recompute meaningful Suggested_Tags for already-built review rows (e.g. a
    glossary loaded from file, or after category edits) — the 'Suggest tags' action.
    Table-level record terms keep their table-level tags."""
    for r in rows or []:
        if not isinstance(r, dict) or r.get("type") == "category":
            continue
        term = str(r.get("Term") or "").strip()
        src = str(r.get("Source_Column") or "").strip()
        if not src and re.search(r"\bRecord$", term):
            continue   # conceptual table term — leave its record/table-level tags
        col = src.split(";")[0].split(".")[-1] if src else term
        cur = [x for x in str(r.get("Suggested_Tags") or "").split(";") if x.strip()]
        is_key = ("identifier" in [c.lower() for c in cur]) or bool(re.search(r"(^|_)id$|_id(_|$)|\bcode\b", col.lower()))
        base = [x for x in cur if x.lower() in {"pii"}]   # preserve an explicit PII flag
        r["Suggested_Tags"] = ";".join(suggest_tags(
            r.get("Category"), r.get("Sensitivity", "LOW"), r.get("PII_Category", ""),
            r.get("Critical_Data_Element", "No"), is_key, base, name=col, term=term))
        lifted = tagdict.lift_sensitivity(r.get("Sensitivity", "LOW"),
                                          r["Suggested_Tags"].split(";"), term)
        if lifted != r.get("Sensitivity"):
            r["Sensitivity"] = lifted
    return rows

def rate_column(confidence=None, has_comment=False, pk=False, fk=False,
                notnull=False, uniqueness=None, completeness=None,
                sensitivity=None, has_term=True, has_definition=True):
    """Suggest a 1-5 star User Rating for a column from scan/profile signals.

    A governance-readiness + data-quality heuristic meant as a starting point a
    steward can override — NOT a substitute for Pentaho Data Quality scoring.
    Signals (all optional, degrades gracefully when profiling wasn't run):
      confidence    suggester confidence in the term mapping (High/Medium/Low)
      has_comment   column documented at the source (DDL COMMENT)
      pk/fk/notnull structural integrity / completeness guarantee
      uniqueness    distinct/non-null ratio from profiling (0-1)
      completeness  non-null/total ratio from profiling (0-1)
      sensitivity   set to a known level (column has been classified)
      has_term/has_definition  governance metadata present
    """
    score = 2.0
    c = str(confidence or "").lower()
    if c == "high":
        score += 1.0
    elif c == "medium":
        score += 0.5
    if has_comment:
        score += 0.5
    if has_definition:
        score += 0.25
    if has_term:
        score += 0.25
    if pk or notnull:
        score += 0.5
    if fk:
        score += 0.25
    if completeness is not None:
        score += (float(completeness) - 0.5)        # +/-0.5 around half-full
    if uniqueness is not None and float(uniqueness) >= 0.9:
        score += 0.5
    if sensitivity and str(sensitivity).upper() in ("LOW", "MEDIUM", "HIGH"):
        score += 0.25
    return max(1, min(5, int(round(score))))


# Default DQ dimension weights (renormalised over whichever dimensions apply
# to a given column). Completeness applies to every column; uniqueness only
# where the column is expected to be unique (keys/identifiers), so a low-
# cardinality enum is not penalised; validity only where a type/pattern was
# detected to conform against.
DQ_WEIGHTS = {"completeness": 0.4, "uniqueness": 0.3, "validity": 0.3}


def quality_score_column(completeness=None, uniqueness=None, validity=None,
                         expect_unique=False, notnull=False, weights=None):
    """Best-practice Data Quality score (0-100) from profiling signals.

    Scores only the dimensions that can be measured for this column and
    renormalises the weights over them, so a column missing a dimension is not
    unfairly dragged down:
      completeness  non-empty / sampled rows (proxy: NOT NULL -> 1.0 if unprofiled)
      uniqueness    distinct / non-null, counted ONLY when uniqueness is expected
                    (primary key or identifier-like) -- otherwise a defect-free
                    low-cardinality column would score badly
      validity      share of values conforming to the detected type/pattern

    Returns an int 0-100, or None when nothing is measurable (so the caller can
    skip writing a qualityScore rather than assert a misleading 0)."""
    w = dict(DQ_WEIGHTS)
    if weights:
        for k in w:
            if weights.get(k) is not None:
                w[k] = float(weights[k])
    dims = []
    comp = completeness
    if comp is None and notnull:
        comp = 1.0
    if comp is not None:
        dims.append((w["completeness"], max(0.0, min(1.0, float(comp)))))
    if expect_unique and uniqueness is not None:
        dims.append((w["uniqueness"], max(0.0, min(1.0, float(uniqueness)))))
    if validity is not None:
        dims.append((w["validity"], max(0.0, min(1.0, float(validity)))))
    wsum = sum(wt for wt, _ in dims)
    if wsum <= 0:
        return None
    score = sum(wt * v for wt, v in dims) / wsum
    return int(round(score * 100))


def rate_document(owner=None, ext=None, sensitivity=None, recent=False, has_term=True):
    """Suggest a 1-5 rating for a FILE/object entity. The column heuristic's
    structural signals (pk/fk/uniqueness/not-null) don't exist for files, so this
    rates curation/governance instead: a steward/owner tag, a usable known format,
    a classified sensitivity, and recency. Conservative midpoint baseline."""
    score = 3.0
    if owner:
        score += 1.0                      # governed: has an owner/steward signal
    known = (ext or "").lower() in (
        "json", "csv", "tsv", "psv", "parquet", "avro", "orc",
        "xml", "xlsx", "xls", "pdf", "docx", "txt")
    if known:
        score += 0.5                      # a format PDC can profile/extract from
    if str(sensitivity or "").upper() in ("MEDIUM", "HIGH"):
        score += 0.5                      # recognised, classified document class
    if recent:
        score += 0.5                      # recently modified -> more likely current
    if not has_term:
        score -= 1.0
    return max(1, min(5, int(round(score))))


def suggest(tables, schema=None):
    schema = schema or os.environ.get("GLOSSARY_SCHEMA", "public")
    """Build suggested glossary rows from scanned tables (term, definition, sensitivity, rating, DQ dims)."""
    rows, seen, out = [], {}, []
    for tname, cols in tables.items():
        category = categorize(tname)
        for c in cols:
            if SKIP.match(c["column"]):
                continue
            name = humanize(c["column"])
            # canonicalize divergent names to one governed term (e.g. "Cust ID" ->
            # "Customer ID"), so instances across tables collapse and merge cleanly.
            _canon = tagdict.canonical_name(name)
            _orig_name = name
            if _canon:
                name = _canon
            pii, sens, tags = classify(c["column"])
            prof = c.get("profile") or {}
            if prof.get("pii"):         pii = prof["pii"]
            if prof.get("sensitivity"): sens = prof["sensitivity"]
            profiled_unique = (prof.get("uniq") or 0) >= 0.95
            is_key = bool(c["pk"] or c["fk"])
            cde = is_cde(name, sens, pii, bool(c["pk"]))
            # CONFIDENCE is an EVIDENCE signal (how sure we are of the term mapping),
            # not a data-quality score. The ladder runs strongest evidence first:
            #   High   - a real DDL comment, a key column, or a profiling hit in data
            #   Medium - the name matched a PII pattern, or weaker profiling evidence
            #   Low    - nothing but the column name to go on (templated)
            # `reason` is surfaced in the UI so the user can see WHY each term scored.
            if c["comment"]:
                conf, reason = "High", "DDL comment used for definition"
            elif is_key:
                conf, reason = "High", "Key column - identity/relationship"
            elif prof.get("confidence") == "High":
                conf, reason = "High", prof.get("reason", "Profiled from data")
            elif pii:
                conf, reason = "Medium", (prof.get("reason") if prof.get("pii") else f"Name matched {pii} pattern")
            elif prof.get("confidence") == "Medium":
                conf, reason = "Medium", prof.get("reason", "Profiled from data")
            else:
                conf, reason = "Low", "Templated from column name"
            if _canon:
                reason = f"{reason} · canonicalized from '{_orig_name}' (dictionary alias)"
            all_tags = suggest_tags(category, sens, pii, "Yes" if cde else "No", is_key, tags, name=c["column"], term=name)
            # lift sensitivity to the highest floor the tags / canonical term imply
            # (ordinal — the dictionary can only tighten a classification, never relax it)
            lifted = tagdict.lift_sensitivity(sens, all_tags, name)
            if lifted != sens:
                sens = lifted
                cde = is_cde(name, sens, pii, bool(c["pk"]))
                all_tags = suggest_tags(category, sens, pii, "Yes" if cde else "No", is_key, tags, name=c["column"], term=name)
            rating = rate_column(confidence=conf, has_comment=bool(c["comment"]),
                                 pk=c["pk"], fk=c["fk"], notnull=c["notnull"],
                                 uniqueness=prof.get("uniq"), sensitivity=sens,
                                 has_term=True, has_definition=True)
            src = f"{schema}.{tname}.{c['column']}"
            # raw DQ dimensions for this physical column (weight-independent, so
            # weights can be tuned later at Apply time without re-scanning)
            expect_unique = bool(c["pk"] or prof.get("kind") in ("identifier", "ssn", "card", "email"))
            qdims = {"c": prof.get("completeness"), "u": prof.get("uniq"),
                     "v": prof.get("valid"), "eu": expect_unique, "nn": bool(c["notnull"])}
            quality = quality_score_column(completeness=qdims["c"], uniqueness=qdims["u"],
                                           validity=qdims["v"], expect_unique=qdims["eu"],
                                           notnull=qdims["nn"])
            rows.append({"Keep": "Y", "Category": category, "Term": name,
                         "Source_Column": src,
                         "Definition": define(c), "Purpose": purpose(c, category, name, pii),
                         "Sensitivity": sens,
                         "PII_Category": pii or "", "Critical_Data_Element": "Yes" if cde else "No",
                         "Abbreviation": _abbrev(name), "Suggested_Tags": ";".join(all_tags),
                         "Suggested_Rating": rating,
                         # per-physical-column rating, so a term mapping to several
                         # columns rates each on its own scan signals (not one shared)
                         "Source_Ratings": {src: rating},
                         # per-column DQ score + the raw dimensions behind it
                         "Suggested_Quality": quality,
                         "Source_Quality_Dims": {src: qdims},
                         # physical key facts (PK/FK + referenced column). PDC's
                         # built-in Is Primary/Foreign Key is harvest-owned metadata
                         # the public API cannot PATCH, so Apply lands these under
                         # attributes.extended and the Registry records them for
                         # the Policy Generator's relationship context.
                         "Source_Keys": ({src: {"pk": bool(c["pk"]), "fk": bool(c["fk"]),
                                                "ref": (f"{c['ref_table']}.{c['ref_col']}"
                                                        if c.get("fk") and c.get("ref_table")
                                                        else None)}}
                                         if (c["pk"] or c["fk"]) else {}),
                         "Status": "Draft", "Confidence": conf, "Suggested_Reason": reason,
                         # scan evidence: the induced value format / reference list —
                         # carried through save + export so the Registry can hand the
                         # Policy Generator a ready-made pattern / dictionary seed
                         "Value_Signature": prof.get("signature", ""),
                         "Value_Pattern": prof.get("pattern", ""),
                         "Enum_Values": ";".join(prof.get("enum", []) or []),
                         "LLM_Enriched": "No"})
    for r in rows:
        key = (r["Category"], r["Term"])
        if key in seen:
            seen[key]["Source_Column"] += "; " + r["Source_Column"]
            # carry each merged column's own rating; representative = best of them
            seen[key].setdefault("Source_Ratings", {}).update(r.get("Source_Ratings", {}))
            seen[key].setdefault("Source_Keys", {}).update(r.get("Source_Keys", {}))
            seen[key]["Suggested_Rating"] = max(seen[key].get("Suggested_Rating", 0),
                                                r.get("Suggested_Rating", 0))
            # carry each merged column's own DQ dimensions
            seen[key].setdefault("Source_Quality_Dims", {}).update(r.get("Source_Quality_Dims", {}))
            for f in ("Value_Signature", "Value_Pattern", "Enum_Values"):
                if not seen[key].get(f) and r.get(f):
                    seen[key][f] = r[f]
            continue
        seen[key] = r
        out.append(r)
    # Add one table-level "record" term per table — created in the glossary so the
    # Steward has it to link, but conceptual (no Source_Column) so the app never
    # auto-links it. Skip any that would collide with an existing (category, term).
    existing = {(r["Category"], r["Term"]) for r in out}
    out += [r for r in table_term_rows(tables, out)
            if (r["Category"], r["Term"]) not in existing]
    return out

# --------------------------------------------------------- DOCUMENT STORE (MinIO/S3)
# Ownership can only be "determined" from the store if it was recorded there.
# We look, in order, at: S3 object tags, x-amz-meta-* user metadata, and bucket
# tags. Keys that look like an owner/steward field are treated as a binding hint
# (resolved later against the people roster by email/username).
OWNER_KEYS = ("steward", "datasteward", "businesssteward", "owner", "custodian", "maintainer")

def _looks_like_owner(key):
    """True when a tag/metadata key looks like an owner/steward field."""
    k = re.sub(r"[^a-z]", "", (key or "").lower())
    return any(o in k for o in OWNER_KEYS)

def _owner_from_pairs(pairs):
    """Pick an owner/steward value out of (key, value) tag/metadata pairs."""
    for k, v in pairs:
        if _looks_like_owner(k) and v:
            return str(v).strip()
    return ""

def _s3_client(cfg):
    """Build an S3 client for a MinIO/S3 endpoint. Raises a clear error if boto3
       is missing or the endpoint is unreachable (callers surface the message)."""
    try:
        import boto3
        from botocore.config import Config
    except Exception:
        raise RuntimeError("boto3 not installed - run: pip install boto3")
    endpoint = (cfg.get("endpoint") or "").strip()
    secure = bool(cfg.get("secure", False))
    if endpoint and not endpoint.startswith(("http://", "https://")):
        endpoint = ("https://" if secure else "http://") + endpoint
    return boto3.client(
        "s3", endpoint_url=endpoint or None,
        aws_access_key_id=cfg.get("access_key") or cfg.get("user") or "",
        aws_secret_access_key=cfg.get("secret_key") or cfg.get("password") or "",
        region_name=cfg.get("region") or "us-east-1",
        config=Config(signature_version="s3v4", s3={"addressing_style": "path"},
                      connect_timeout=8, read_timeout=15, retries={"max_attempts": 2}))

def test_minio(cfg):
    """Verify a MinIO/S3 connection without a full scan: {ok, message, objects, tagging}."""
    bucket = (cfg.get("bucket") or "").strip()
    if not bucket:
        return {"ok": False, "message": "No bucket specified"}
    try:
        s3 = _s3_client(cfg)
    except Exception as e:
        return {"ok": False, "needs_driver": "boto3" in str(e), "message": str(e)}
    try:
        resp = s3.list_objects_v2(Bucket=bucket, MaxKeys=1)
        n = resp.get("KeyCount", 0)
        tagging = False
        if resp.get("Contents"):
            try:
                s3.get_object_tagging(Bucket=bucket, Key=resp["Contents"][0]["Key"])
                tagging = True
            except Exception:
                tagging = False
        return {"ok": True, "message": "Bucket reachable", "objects": n, "tagging": tagging}
    except Exception as ex:
        msg = f"Connection failed: {ex}"
        s = str(ex)
        if "WRONG_VERSION_NUMBER" in s:
            msg += (" — the endpoint answered plain HTTP to a TLS handshake: this "
                    "port has no TLS. Use http:// in the endpoint and untick HTTPS "
                    "(MinIO on :9000 is usually plain HTTP in the lab).")
        elif "record layer failure" in s or "UNEXPECTED_RECORD" in s:
            msg += (" — scheme mismatch between the endpoint URL and the TLS toggle; "
                    "make http/https and the HTTPS tick agree.")
        return {"ok": False, "message": msg}

_TEXT_EXTS = {"txt", "csv", "tsv", "json", "jsonl", "ndjson", "xml", "md", "log",
              "yaml", "yml", "html", "htm", "sql", "py", "sh", "conf", "ini", "properties"}
_IMAGE_EXTS = {"png", "jpg", "jpeg", "gif", "webp", "bmp", "svg"}
_CTYPES = {"pdf": "application/pdf", "png": "image/png", "jpg": "image/jpeg",
           "jpeg": "image/jpeg", "gif": "image/gif", "webp": "image/webp",
           "bmp": "image/bmp", "svg": "image/svg+xml",
           "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
           "txt": "text/plain", "csv": "text/csv", "json": "application/json",
           "html": "text/html", "htm": "text/html", "xml": "application/xml"}


def _ext_of(key):
    """Lowercased extension of an object key (last path segment)."""
    leaf = key.rsplit("/", 1)[-1]
    return leaf.rsplit(".", 1)[-1].lower() if "." in leaf else ""


def _guess_ctype(key, fallback=""):
    """Content-type from extension (authoritative for known viewable types), so a PDF
       streamed from MinIO renders even when the store reports octet-stream."""
    return _CTYPES.get(_ext_of(key)) or fallback or "application/octet-stream"


def list_objects(cfg, prefix="", max_keys=2000):
    """List ONE level of a bucket (S3 delimiter='/') for the file browser:
       subfolders (CommonPrefixes) + files at this prefix with size/modified/ext.
       `prefix` is relative to the connection's configured base prefix."""
    s3 = _s3_client(cfg)
    bucket = (cfg.get("bucket") or "").strip()
    base = (cfg.get("prefix") or "").strip().strip("/")
    bpref = base + "/" if base else ""
    rel = (prefix or "").strip().strip("/")
    full = bpref + (rel + "/" if rel else "")
    folders, files, count, truncated = [], [], 0, False
    for page in s3.get_paginator("list_objects_v2").paginate(
            Bucket=bucket, Prefix=full, Delimiter="/"):
        for cp in page.get("CommonPrefixes", []):
            p = cp.get("Prefix", "")
            name = p[len(full):].rstrip("/")
            if name:
                folders.append({"name": name, "prefix": p[len(bpref):].rstrip("/")})
        for obj in page.get("Contents", []):
            key = obj["Key"]
            if key.endswith("/") or key == full:
                continue
            name = key[len(full):]
            if "/" in name:
                continue
            ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
            mod = obj.get("LastModified")
            files.append({"name": name, "key": key, "size": obj.get("Size", 0),
                          "modified": mod.isoformat() if mod else "", "ext": ext})
            count += 1
            if count >= max_keys:
                truncated = True
                break
        if truncated:
            break
    folders.sort(key=lambda x: x["name"].lower())
    files.sort(key=lambda x: x["name"].lower())
    total_bytes = sum(f["size"] for f in files)
    return {"bucket": bucket, "base_prefix": base, "prefix": rel,
            "folders": folders, "files": files, "total_bytes": total_bytes,
            "file_count": len(files), "folder_count": len(folders), "truncated": truncated}


def object_detail(cfg, key, max_preview=6000):
    """Metadata + tags (+ a short text preview for text-like files) for one object."""
    s3 = _s3_client(cfg)
    bucket = (cfg.get("bucket") or "").strip()
    out = {"key": key, "name": key.rsplit("/", 1)[-1], "tags": [], "metadata": {},
           "preview": None, "preview_truncated": False}
    try:
        h = s3.head_object(Bucket=bucket, Key=key)
        out["size"] = h.get("ContentLength", 0)
        out["content_type"] = h.get("ContentType", "")
        out["modified"] = h.get("LastModified").isoformat() if h.get("LastModified") else ""
        out["metadata"] = dict(h.get("Metadata", {}) or {})
    except Exception as e:
        out["error"] = str(e)
        return out
    try:
        ts = s3.get_object_tagging(Bucket=bucket, Key=key).get("TagSet", [])
        out["tags"] = [{"key": t["Key"], "value": t["Value"]} for t in ts]
    except Exception:
        out["tags"] = []
    ext = key.rsplit(".", 1)[-1].lower() if "." in key.rsplit("/", 1)[-1] else ""
    out["ext"] = ext
    # classify how the UI should show it
    if ext == "pdf":
        out["preview_kind"] = "pdf"
    elif ext == "docx":
        out["preview_kind"] = "docx"
    elif ext in _IMAGE_EXTS:
        out["preview_kind"] = "image"
    elif ext in _TEXT_EXTS:
        out["preview_kind"] = "text"
    else:
        out["preview_kind"] = "none"

    if ext in _TEXT_EXTS and out.get("size", 0) > 0:
        raw = _get_object_bytes(s3, bucket, key, max_preview)
        if raw:
            try:
                out["preview"] = raw.decode("utf-8", errors="replace")
            except Exception:
                out["preview"] = None
            out["preview_truncated"] = out.get("size", 0) > max_preview
    elif ext == "docx" and 0 < out.get("size", 0) <= _DOCX_MAX:
        # Word docs can't render natively, so convert to HTML server-side
        try:
            data = s3.get_object(Bucket=bucket, Key=key)["Body"].read()
            html, how = docx_to_html(data)
            out["html"] = html
            out["docx_renderer"] = how
        except Exception as e:
            out["html"] = None
            out["docx_error"] = str(e)
    return out


_VIEW_MAX = 25 * 1024 * 1024   # cap inline viewing at 25 MB
_DOCX_MAX = 8 * 1024 * 1024    # convert Word docs up to 8 MB


def get_object_bytes_full(cfg, key, max_bytes=_VIEW_MAX):
    """Fetch a whole object for inline viewing (PDF/image), with its content-type.
       Refuses files over max_bytes so a huge object can't be pulled into a view."""
    s3 = _s3_client(cfg)
    bucket = (cfg.get("bucket") or "").strip()
    h = s3.head_object(Bucket=bucket, Key=key)
    size = h.get("ContentLength", 0)
    if size and size > max_bytes:
        raise ValueError(f"File is {size:,} bytes — too large to view inline "
                         f"(limit {max_bytes // (1024*1024)} MB). Use Download instead.")
    ctype = _guess_ctype(key, h.get("ContentType", ""))
    data = s3.get_object(Bucket=bucket, Key=key)["Body"].read()
    return data, ctype


def docx_to_html(data):
    """Render a .docx to HTML. Prefers mammoth (keeps headings/lists/bold); falls back
       to python-docx (paragraph text); returns (html_or_None, renderer_name)."""
    import io
    try:
        import mammoth
        res = mammoth.convert_to_html(io.BytesIO(data))
        return res.value or "<p><em>(empty document)</em></p>", "mammoth"
    except ImportError:
        pass
    except Exception as e:
        return None, f"mammoth-error: {e}"
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if not t:
                continue
            style = (p.style.name or "").lower() if p.style else ""
            if "heading 1" in style or style == "title":
                parts.append(f"<h2>{_html_escape(t)}</h2>")
            elif "heading" in style:
                parts.append(f"<h3>{_html_escape(t)}</h3>")
            elif "list" in style:
                parts.append(f"<li>{_html_escape(t)}</li>")
            else:
                parts.append(f"<p>{_html_escape(t)}</p>")
        for tbl in getattr(doc, "tables", []):
            rows = []
            for r in tbl.rows:
                cells = "".join(f"<td>{_html_escape(c.text.strip())}</td>" for c in r.cells)
                rows.append(f"<tr>{cells}</tr>")
            if rows:
                parts.append("<table border='1' cellspacing='0' cellpadding='4'>"
                             + "".join(rows) + "</table>")
        return ("\n".join(parts) or "<p><em>(empty document)</em></p>", "python-docx")
    except ImportError:
        return None, "no-renderer"
    except Exception as e:
        return None, f"docx-error: {e}"


def _html_escape(s):
    return (str(s).replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;"))


def harvest_minio(cfg, max_objects=20000, tag_sample=600):
    """Walk a bucket, roll objects into top-level folders, and sniff ownership.
       Returns (folders_dict, ownership_report, scanned_counts)."""
    s3 = _s3_client(cfg)
    bucket = (cfg.get("bucket") or "").strip()
    prefix = (cfg.get("prefix") or "").lstrip("/")
    folders, n, tagged = {}, 0, 0
    ownership = {"bucket_owner": "", "by_folder": {}, "signals": [], "tags_sampled": 0}

    try:
        bt = s3.get_bucket_tagging(Bucket=bucket).get("TagSet", [])
        ownership["bucket_owner"] = _owner_from_pairs((t["Key"], t["Value"]) for t in bt)
    except Exception:
        pass

    for page in s3.get_paginator("list_objects_v2").paginate(Bucket=bucket, Prefix=prefix):
        for obj in page.get("Contents", []):
            key = obj["Key"]
            if key.endswith("/"):
                continue
            rel = key[len(prefix):].lstrip("/") if prefix else key
            top = rel.split("/")[0] if "/" in rel else "(root)"
            fpref = "/".join([p for p in [prefix.rstrip("/"), (top if top != "(root)" else "")] if p])
            f = folders.setdefault(top, {"name": top, "prefix": fpref, "count": 0,
                                         "bytes": 0, "exts": {}, "samples": [], "owners": {}})
            f["count"] += 1
            f["bytes"] += obj.get("Size", 0)
            base = key.rsplit("/", 1)[-1]
            if "." in base:
                ext = base.rsplit(".", 1)[-1].lower()
                f["exts"][ext] = f["exts"].get(ext, 0) + 1
            if len(f["samples"]) < 3:
                f["samples"].append(rel)
            n += 1
            if tagged < tag_sample:
                owner = ""
                try:
                    ts = s3.get_object_tagging(Bucket=bucket, Key=key).get("TagSet", [])
                    owner = _owner_from_pairs((t["Key"], t["Value"]) for t in ts)
                except Exception:
                    pass
                if not owner:
                    try:
                        meta = s3.head_object(Bucket=bucket, Key=key).get("Metadata", {})
                        owner = _owner_from_pairs(meta.items())
                    except Exception:
                        pass
                if owner:
                    f["owners"][owner] = f["owners"].get(owner, 0) + 1
                tagged += 1
            if n >= max_objects:
                break
        if n >= max_objects:
            break

    ownership["tags_sampled"] = tagged
    for top, f in folders.items():
        if f["owners"]:
            f["owner"] = max(f["owners"], key=f["owners"].get)
            ownership["by_folder"][top] = f["owner"]
    if ownership["bucket_owner"]:
        ownership["signals"].append(f"Bucket tag owner: {ownership['bucket_owner']}")
    for top, ow in ownership["by_folder"].items():
        ownership["signals"].append(f"Folder '{top}' owner tag/metadata: {ow}")
    if not ownership["signals"]:
        ownership["signals"].append(
            f"No owner/steward signal found in object tags or x-amz-meta-* metadata "
            f"({tagged} object(s) sampled). Assign the steward manually.")
    return folders, ownership, {"objects": n, "folders": len(folders)}

# ----------------------------------------------------- DOCUMENT DATA-QUALITY PROFILER
# These helpers let the app compute a real Data-Quality score for object-store files
# by reading their CONTENT, instead of leaving the Data Quality input blank for PDC
# to fill later. The output is a dims dict {c,u,v,eu,nn} that is byte-for-byte
# compatible with the Source_Quality_Dims that SQL columns produce, so a document
# rides the SAME weighted scorer (quality_score_column) and the SAME apply pipeline
# (data_element_links -> links_to_api_json -> features.qualityScore -> PATCH).

# File formats we can profile from bytes. Anything else (pdf, docx, images, binary)
# returns None and is left to PDC's own document profiling.
_DQ_TEXT_EXTS  = {"txt", "md", "log", "rtf"}      # free text / prose / logs
_DQ_JSON_EXTS  = {"json"}                          # one whole-document JSON value
_DQ_JSONL_EXTS = {"jsonl", "ndjson"}              # one JSON record per line
_DQ_DELIM_EXTS = {"csv", "tsv", "psv"}            # delimited record sets
_DQ_XML_EXTS   = {"xml"}                          # one whole-document XML tree
_DQ_PROFILABLE = (_DQ_TEXT_EXTS | _DQ_JSON_EXTS | _DQ_JSONL_EXTS
                  | _DQ_DELIM_EXTS | _DQ_XML_EXTS)
# Line-oriented formats can be profiled from a truncated HEAD of a large file (we drop
# the last, possibly-incomplete line). Whole-document formats (json/xml) cannot be
# truncated without breaking the parse, so large ones are skipped (-> defer to PDC).
_DQ_LINE_EXTS = _DQ_TEXT_EXTS | _DQ_JSONL_EXTS | _DQ_DELIM_EXTS


def _completeness_of_records(records):
    """completeness = non-empty values / total values across a list of flat records.
       A value is 'empty' if it is None or a blank/whitespace-only string. Nested
       dict/list values count as present (non-empty) -- we only penalise true gaps.
       Returns a 0.0-1.0 fraction, or None when there is nothing to measure."""
    total = filled = 0
    for rec in records:
        if isinstance(rec, dict):
            for v in rec.values():
                total += 1
                if v is None:
                    continue
                if isinstance(v, str) and not v.strip():
                    continue
                filled += 1
        else:
            # a bare (non-dict) record counts as a single cell
            total += 1
            if rec not in (None, "") and not (isinstance(rec, str) and not rec.strip()):
                filled += 1
    return (filled / total) if total else None


def _uniqueness_of(rows):
    """uniqueness = distinct / total over row signatures. A high value means few
       duplicate records, which is a data-quality positive. JSON-encodes each row
       with sorted keys so dict order doesn't create false 'differences'."""
    if not rows:
        return None
    sigs = [json.dumps(r, sort_keys=True, default=str) for r in rows]
    return len(set(sigs)) / len(sigs)


def profile_document_object(content, ext):
    """Compute Data-Quality dimensions for ONE object's content.

       Returns a dims dict shaped exactly like a column's Source_Quality_Dims entry:
         c  completeness  non-empty values / total values
         v  validity      well-formedness: parses cleanly / decodes as UTF-8
         u  uniqueness    distinct records / total (duplicate detection)
         eu expect_unique whether uniqueness should count toward the score. True only
                          for record-set files (JSON array / JSONL / delimited),
                          where duplicate rows are a genuine defect; False for a
                          single object or free text, where 'uniqueness' is meaningless
         nn notnull       always False here -- we measure completeness directly rather
                          than inferring it from a NOT NULL constraint
       Returns None when the extension isn't content-profilable (pdf/docx/binary),
       so the caller leaves the Data Quality input to PDC."""
    ext = (ext or "").lower()
    if not content or ext not in _DQ_PROFILABLE:
        return None

    # Decode to text. For a format that is supposed to be text, undecodable bytes are
    # themselves a validity defect, so we remember whether the UTF-8 decode succeeded.
    try:
        text = content.decode("utf-8")
        decoded = True
    except Exception:
        try:
            text = content.decode("latin-1")   # last-resort so we can still inspect
        except Exception:
            return None
        decoded = False

    # ---- JSON: a single whole-document value (object, array, or scalar) ----------
    if ext in _DQ_JSON_EXTS:
        try:
            obj = json.loads(text)
        except Exception:
            # malformed JSON -> validity 0; nothing else is trustworthy to measure
            return {"c": None, "u": None, "v": 0.0, "eu": False, "nn": False}
        if isinstance(obj, list):                       # array of records
            return {"c": _completeness_of_records(obj), "u": _uniqueness_of(obj),
                    "v": 1.0, "eu": True, "nn": False}
        if isinstance(obj, dict):                       # one record
            return {"c": _completeness_of_records([obj]), "u": None,
                    "v": 1.0, "eu": False, "nn": False}
        return {"c": 1.0 if obj not in (None, "") else 0.0, "u": None,
                "v": 1.0, "eu": False, "nn": False}     # bare scalar, still well-formed

    # ---- JSONL / NDJSON: one JSON record per line --------------------------------
    if ext in _DQ_JSONL_EXTS:
        recs, bad = [], 0
        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                recs.append(json.loads(line))
            except Exception:
                bad += 1                                # a line that won't parse = defect
        total = len(recs) + bad
        validity = (len(recs) / total) if total else 0.0
        return {"c": _completeness_of_records(recs), "u": _uniqueness_of(recs),
                "v": validity, "eu": True, "nn": False}

    # ---- delimited record sets: CSV / TSV / PSV ----------------------------------
    if ext in _DQ_DELIM_EXTS:
        import csv, io
        delim = {"csv": ",", "tsv": "\t", "psv": "|"}[ext]
        rows = [r for r in csv.reader(io.StringIO(text), delimiter=delim) if r]
        if not rows:
            return {"c": None, "u": None, "v": 1.0 if decoded else 0.0,
                    "eu": False, "nn": False}
        header, body = rows[0], rows[1:]
        ncol = len(header)
        # validity = share of body rows whose column count matches the header; a
        # ragged row (too few/many fields) is a structural quality defect.
        validity = (sum(1 for r in body if len(r) == ncol) / len(body)) if body else 1.0
        # completeness = non-empty cells / total cells across the body rows.
        cells = filled = 0
        for r in body:
            for v in r:
                cells += 1
                if str(v).strip():
                    filled += 1
        completeness = (filled / cells) if cells else None
        # uniqueness = distinct body rows / total -> duplicate-record detection.
        uniqueness = _uniqueness_of([tuple(r) for r in body]) if body else None
        return {"c": completeness, "u": uniqueness, "v": validity,
                "eu": True, "nn": False}

    # ---- XML: one whole-document tree --------------------------------------------
    if ext in _DQ_XML_EXTS:
        try:
            import xml.etree.ElementTree as ET
            root = ET.fromstring(text)
        except Exception:
            return {"c": None, "u": None, "v": 0.0, "eu": False, "nn": False}
        # completeness = non-empty attribute values + non-empty leaf text / all slots.
        total = filled = 0
        for el in root.iter():
            for av in el.attrib.values():
                total += 1
                if str(av).strip():
                    filled += 1
            if not list(el):                            # leaf element -> has text slot
                total += 1
                if (el.text or "").strip():
                    filled += 1
        completeness = (filled / total) if total else None
        return {"c": completeness, "u": None, "v": 1.0, "eu": False, "nn": False}

    # ---- plain text / markdown / logs --------------------------------------------
    # completeness = non-blank lines / total lines (an all-whitespace file is poor);
    # validity = decoded cleanly as UTF-8; uniqueness is not meaningful for prose.
    lines = text.splitlines()
    nonblank = sum(1 for ln in lines if ln.strip())
    completeness = (nonblank / len(lines)) if lines else (1.0 if text.strip() else 0.0)
    return {"c": completeness, "u": None, "v": 1.0 if decoded else 0.0,
            "eu": False, "nn": False}


def _get_object_bytes(s3, bucket, key, max_bytes):
    """Fetch up to max_bytes of an object for profiling using a ranged GET, so a large
       file never pulls in whole. Returns bytes (possibly a truncated head) or b''.
       Sampling a head is acceptable for DQ the same way column profiling samples rows
       rather than scanning every value."""
    try:
        resp = s3.get_object(Bucket=bucket, Key=key, Range=f"bytes=0-{max_bytes - 1}")
        return resp["Body"].read()
    except Exception:
        # some stores reject Range on tiny objects; fall back to a capped full GET
        try:
            return s3.get_object(Bucket=bucket, Key=key)["Body"].read(max_bytes)
        except Exception:
            return b""


def harvest_files(cfg, max_objects=20000, owner_sample=600,
                  profile_dq=False, dq_max_bytes=5_000_000, dq_sample=800):
    """Enumerate INDIVIDUAL objects (leaf files) in a bucket, retaining each key so
       metadata can be applied per file (harvest_minio rolls these up into folders and
       discards the keys). Honours the same include/exclude globs as discover_documents
       and samples owner tags/metadata up to owner_sample objects. Returns a list of
       file dicts: {key, rel, bucket, folder, base, ext, size, owner, recent}.

       When profile_dq is set, each content-profilable file (csv/tsv/psv/json/jsonl/
       xml/txt/md/log) up to dq_sample files is also READ (a head of up to dq_max_bytes)
       and scored for Data Quality; the resulting dims land on the file dict as 'qdims',
       which suggest_document_files turns into Source_Quality_Dims so the file's
       qualityScore is computed and PATCHed exactly like a SQL column's."""
    s3 = _s3_client(cfg)
    bucket = (cfg.get("bucket") or "").strip()
    prefix = (cfg.get("prefix") or "").lstrip("/")
    include = _doc_patterns(cfg.get("include"))
    exclude = _doc_patterns(cfg.get("exclude"))
    import datetime
    now = datetime.datetime.now(datetime.timezone.utc)
    files, n, sniffed, dq_done = [], 0, 0, 0
    for page in s3.get_paginator("list_objects_v2").paginate(Bucket=bucket, Prefix=prefix):
        for obj in page.get("Contents", []):
            key = obj["Key"]
            if key.endswith("/"):
                continue
            rel = key[len(prefix):].lstrip("/") if prefix else key
            base = key.rsplit("/", 1)[-1]
            if include and not _doc_match(rel, base, include):
                continue
            if exclude and _doc_match(rel, base, exclude):
                continue
            top = rel.split("/")[0] if "/" in rel else "(root)"
            ext = base.rsplit(".", 1)[-1].lower() if "." in base else ""
            lm = obj.get("LastModified")
            recent = False
            if lm is not None:
                try:
                    recent = (now - lm).days <= 90
                except Exception:
                    recent = False
            owner = ""
            if sniffed < owner_sample:
                try:
                    ts = s3.get_object_tagging(Bucket=bucket, Key=key).get("TagSet", [])
                    owner = _owner_from_pairs((t["Key"], t["Value"]) for t in ts)
                except Exception:
                    pass
                if not owner:
                    try:
                        meta = s3.head_object(Bucket=bucket, Key=key).get("Metadata", {})
                        owner = _owner_from_pairs(meta.items())
                    except Exception:
                        pass
                sniffed += 1
            rec_file = {"key": key, "rel": rel, "bucket": bucket, "folder": top,
                        "base": base, "ext": ext, "size": obj.get("Size", 0),
                        "owner": owner, "recent": recent}
            # Optional: read the object and compute a Data-Quality score from content.
            # Bounded two ways: only the first dq_sample files are profiled, and only a
            # head of up to dq_max_bytes is read per file. Whole-document formats
            # (json/xml) are skipped when larger than the cap (truncation would break
            # the parse); line-oriented formats are profiled from a head with the last
            # partial line dropped.
            if profile_dq and dq_done < dq_sample and ext in _DQ_PROFILABLE:
                size = obj.get("Size", 0) or 0
                if size <= dq_max_bytes or ext in _DQ_LINE_EXTS:
                    raw = _get_object_bytes(s3, bucket, key, dq_max_bytes)
                    if raw and size > dq_max_bytes and ext in _DQ_LINE_EXTS:
                        raw = raw.rsplit(b"\n", 1)[0]      # drop truncated final line
                    qd = profile_document_object(raw, ext)
                    if qd:
                        rec_file["qdims"] = qd
                    dq_done += 1
            files.append(rec_file)
            n += 1
            if n >= max_objects:
                break
        if n >= max_objects:
            break
    return files

def _doc_patterns(s):
    """Parse a comma/newline-separated glob string into a lowercased pattern list.
       e.g. '*.md, inspections/*' -> ['*.md', 'inspections/*']."""
    if not s:
        return []
    return [p.strip().lower() for p in re.split(r"[,\n]+", str(s)) if p.strip()]


def _doc_match(rel, base, pats):
    """True if the object matches any glob pattern, tested against both its
       relative key (so 'inspections/*' works) and its basename (so '*.md' works)."""
    import fnmatch
    rl, bl = rel.lower(), base.lower()
    return any(fnmatch.fnmatch(bl, p) or fnmatch.fnmatch(rl, p) for p in pats)


def discover_documents(cfg, max_objects=50000, top_n=8):
    """High-level discovery of a bucket's contents: file counts, total size,
       breakdown by file type and by folder, plus largest and newest objects.

       Optional cfg['include'] / cfg['exclude'] are comma/newline-separated glob
       patterns. An object is kept when it matches an include pattern (or none are
       given) and matches no exclude pattern. Patterns test the basename and the
       relative key, so '*.md' drops all Markdown and 'inspections/*' scopes a
       folder. Filtered objects are counted but excluded from every roll-up."""
    import heapq
    s3 = _s3_client(cfg)
    bucket = (cfg.get("bucket") or "").strip()
    prefix = (cfg.get("prefix") or "").lstrip("/")
    include = _doc_patterns(cfg.get("include"))
    exclude = _doc_patterns(cfg.get("exclude"))
    files = total = filtered = 0
    by_type, by_folder = {}, {}
    largest, newest = [], []
    for page in s3.get_paginator("list_objects_v2").paginate(Bucket=bucket, Prefix=prefix):
        for obj in page.get("Contents", []):
            key = obj["Key"]
            if key.endswith("/"):
                continue
            size = obj.get("Size", 0) or 0
            lm = obj.get("LastModified")
            rel = key[len(prefix):].lstrip("/") if prefix else key
            base = key.rsplit("/", 1)[-1]
            # include/exclude glob filtering (before the object counts anywhere)
            if include and not _doc_match(rel, base, include):
                filtered += 1
                continue
            if exclude and _doc_match(rel, base, exclude):
                filtered += 1
                continue
            ext = base.rsplit(".", 1)[-1].lower() if "." in base else "(none)"
            top = rel.split("/")[0] if "/" in rel else "(root)"
            files += 1; total += size
            t = by_type.setdefault(ext, {"ext": ext, "count": 0, "bytes": 0})
            t["count"] += 1; t["bytes"] += size
            fo = by_folder.setdefault(top, {"name": top, "count": 0, "bytes": 0})
            fo["count"] += 1; fo["bytes"] += size
            heapq.heappush(largest, (size, key))
            if len(largest) > top_n:
                heapq.heappop(largest)
            if lm is not None:
                heapq.heappush(newest, (lm.timestamp(), key, lm.isoformat()))
                if len(newest) > top_n:
                    heapq.heappop(newest)
            if files >= max_objects:
                break
        if files >= max_objects:
            break
    summary = {"files": files, "bytes": total, "folders": len(by_folder),
               "types": len(by_type), "avg_bytes": round(total / files) if files else 0,
               "filtered": filtered}
    return {
        "bucket": bucket, "prefix": prefix, "summary": summary,
        "include": cfg.get("include") or "", "exclude": cfg.get("exclude") or "",
        "by_type": sorted(by_type.values(), key=lambda x: x["count"], reverse=True),
        "by_folder": sorted(by_folder.values(), key=lambda x: x["bytes"], reverse=True),
        "largest": [{"key": k, "bytes": s} for s, k in sorted(largest, reverse=True)],
        "newest": [{"key": k, "modified": iso} for _, k, iso in sorted(newest, reverse=True)],
    }

DOC_RULES = [
    (r"complian|legal|audit|regulat|consent|privacy", "HIGH", ["document", "compliance"], True),
    (r"customer|account|billing|invoice|payment|financ", "MEDIUM", ["document", "pii"], False),
    (r"qualit|lab|test|sampl|inspect", "MEDIUM", ["document"], False),
    (r"public|report|notice|brochure|template", "LOW", ["document"], False),
]

def _doc_classify(folder):
    """Classify a document folder into (sensitivity, tags, is_critical_data_element)."""
    fl = folder.lower()
    for pat, sens, tags, cde in DOC_RULES:
        if re.search(pat, fl):
            return sens, list(tags), cde
    return "LOW", ["document"], False

def suggest_documents(folders, bucket="documents"):
    """Turn harvested document folders into review rows under 'Records & Documents'.
       Carries an Owner_Hint when the store recorded an owner/steward."""
    rows = []
    for top, f in folders.items():
        name = humanize(top.replace("-", " ").replace("_", " ")) if top != "(root)" else "Bucket Root"
        sens, tags, cde = _doc_classify(top)
        exts = ", ".join(sorted(f["exts"])[:5]) if f["exts"] else "mixed"
        owner = f.get("owner", "")
        defn = (f"Document folder containing {f['count']} object(s) ({exts}) "
                f"in the {bucket} object store.")
        purp = (f"Holds {top.replace('-', ' ').replace('_', ' ')} documents for reference, "
                f"audit, and compliance." if top != "(root)"
                else "Holds supporting documents for reference and compliance.")
        conf = "High" if owner else ("Medium" if f["exts"] else "Low")
        reason = f"Owner tag/metadata: {owner}" if owner else "Derived from object-store folder"
        doc_tags = suggest_tags("Records & Documents", sens, "", "Yes" if cde else "No", False, tags, name=name, term=name)
        rows.append({"Keep": "Y", "Category": "Records & Documents", "Term": name,
                     "Source_Column": f"{bucket}/{f['prefix']}".rstrip("/"),
                     "Definition": defn, "Purpose": purp, "Sensitivity": sens, "PII_Category": "",
                     "Critical_Data_Element": "Yes" if cde else "No", "Abbreviation": "",
                     "Suggested_Tags": ";".join(doc_tags), "Status": "Draft", "Confidence": conf,
                     "Suggested_Reason": reason, "LLM_Enriched": "No", "Owner_Hint": owner})
    return rows

def suggest_document_files(files, bucket="documents"):
    """Per-FILE review rows (leaf objects), so metadata lands on the individual files
       you see in PDC rather than only the folder. Each file inherits its folder's
       business term + sensitivity classification and carries its own document rating
       (keyed by the full object path in Source_Ratings, so it survives term dedup).

       Data Quality: when harvest_files was run with profile_dq (so a file carries
       'qdims' from reading its content), those dimensions are attached here as
       Source_Quality_Dims -- exactly like a SQL column -- so the file gets a computed
       qualityScore through the normal weighted pipeline. Files without 'qdims'
       (un-profilable formats like PDF/DOCX, or DQ left off) attach no dims and defer
       the Data Quality input to PDC's own file profiling; the app still sets the other
       three Trust-Score inputs (term, verified lineage, rating) + sensitivity."""
    rows, seen, out = [], {}, []
    for f in files:
        folder = f.get("folder") or "(root)"
        base = f.get("base") or (f.get("key", "").rsplit("/", 1)[-1])
        if not base:
            continue
        bkt = f.get("bucket") or bucket
        sens, tags, cde = _doc_classify(folder)
        term = (humanize(folder.replace("-", " ").replace("_", " "))
                if folder != "(root)" else "Bucket Root")
        src = f"{bkt}/{folder}/{base}" if folder != "(root)" else f"{bkt}/{base}"
        rating = rate_document(owner=f.get("owner"), ext=f.get("ext"),
                               sensitivity=sens, recent=f.get("recent"))
        doc_tags = suggest_tags("Records & Documents", sens, "", "Yes" if cde else "No", False, tags, name=term, term=term)
        row = {"Keep": "Y", "Category": "Records & Documents", "Term": term,
               "Source_Column": src,
               "Definition": f"Object '{base}' in the {bkt}/{folder} object store.",
               "Purpose": f"Holds {term} data for reference, audit, and compliance.",
               "Sensitivity": sens, "PII_Category": "",
               "Critical_Data_Element": "Yes" if cde else "No", "Abbreviation": "",
               "Suggested_Tags": ";".join(doc_tags),
               "Suggested_Rating": rating, "Source_Ratings": {src: rating},
               "Status": "Draft", "Confidence": "High" if f.get("owner") else "Medium",
               "Suggested_Reason": (f"Owner tag/metadata: {f['owner']}" if f.get("owner")
                                    else "Leaf object in classified folder"),
               "LLM_Enriched": "No", "Owner_Hint": f.get("owner", "")}
        # If the object's content was profiled, carry the dimensions (keyed by the same
        # Source_Column path) so data_element_links computes a qualityScore for it.
        qd = f.get("qdims")
        if qd:
            row["Source_Quality_Dims"] = {src: qd}
            row["Suggested_Quality"] = quality_score_column(
                completeness=qd.get("c"), uniqueness=qd.get("u"), validity=qd.get("v"),
                expect_unique=qd.get("eu"), notnull=qd.get("nn"))
        rows.append(row)
    # dedup by (Category, Term): files sharing a folder term merge into one row,
    # each keeping its own Source_Column + per-file rating + per-file DQ dims
    for r in rows:
        key = (r["Category"], r["Term"])
        if key in seen:
            seen[key]["Source_Column"] += "; " + r["Source_Column"]
            seen[key].setdefault("Source_Ratings", {}).update(r.get("Source_Ratings", {}))
            seen[key]["Suggested_Rating"] = max(seen[key].get("Suggested_Rating", 0),
                                                r.get("Suggested_Rating", 0))
            if r.get("Source_Quality_Dims"):
                seen[key].setdefault("Source_Quality_Dims", {}).update(r["Source_Quality_Dims"])
            continue
        seen[key] = r
        out.append(r)
    return out

# ----------------------------------------------------------------- ENHANCE
def _plain_lex(s):
    """Extract plain text from a Lexical JSON string (or pass through plain text)."""
    if not s:
        return ""
    try:
        o = json.loads(s)
        def walk(n):
            """Recursively return the first text node found in a Lexical-JSON tree."""
            if isinstance(n, dict):
                if n.get("type") == "text":
                    return n.get("text", "")
                for c in n.get("children", []) or []:
                    r = walk(c)
                    if r:
                        return r
            return ""
        return walk(o.get("root", {})) or str(s)
    except Exception:
        return str(s)

def parse_glossary(jsonl_text):
    """Index an exported glossary (JSONL) by term name for enhancement."""
    catname, terms, raw = {}, {}, []
    gname = None
    for line in (jsonl_text or "").splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            r = json.loads(line)
        except Exception:
            continue
        t = r.get("type")
        if t == "glossary":
            gname = r.get("name")
        elif t == "category":
            catname[r.get("_id")] = r.get("name")
        elif t == "term":
            raw.append(r)
    for r in raw:
        a = r.get("attributes", {}); info = a.get("info", {}); feat = a.get("features", {})
        terms[r["name"].strip().lower()] = {
            "Term": r["name"], "Category": catname.get(r.get("parentId"), "Uncategorized"),
            "definition": _plain_lex(info.get("definition")),
            "purpose": _plain_lex(info.get("purpose")),
            "sensitivity": feat.get("sensitivity"), "classification": info.get("classification"),
            "cde": feat.get("isCriticalDataElement"),
            "tags": [str(x.get("name")).strip().lower()
                     for x in a.get("tags", []) if x.get("name")]}
    return {"name": gname, "categories": list(catname.values()), "terms": terms}

def enhance_from_glossary(rows, jsonl_text, append_missing=True):
    """Overlay an existing glossary's real metadata onto matched scanned rows,
       and optionally append export terms the scan didn't produce."""
    g = parse_glossary(jsonl_text)
    idx = g["terms"]
    matched, present = 0, {r["Term"].strip().lower() for r in rows}
    for r in rows:
        e = idx.get(r["Term"].strip().lower())
        if not e:
            continue
        matched += 1
        if e["definition"]:  r["Definition"] = e["definition"]
        if e["purpose"]:     r["Purpose"] = e["purpose"]
        if e["sensitivity"]: r["Sensitivity"] = e["sensitivity"]
        if e["cde"] is not None:
            r["Critical_Data_Element"] = "Yes" if e["cde"] else "No"
        if e["tags"]:
            cur = r.get("Suggested_Tags", "").split(";") if r.get("Suggested_Tags") else []
            r["Suggested_Tags"] = ";".join(dict.fromkeys(cur + e["tags"]))
        r["Confidence"] = "High"
        r["Suggested_Reason"] = "Matched existing glossary term"
    added = []
    if append_missing:
        for key, e in idx.items():
            if key in present:
                continue
            sens = e["sensitivity"] or "LOW"
            added.append({"Keep": "Y", "Category": e["Category"], "Term": e["Term"],
                          "Source_Column": f"glossary:{g.get('name') or 'existing'}",
                          "Definition": e["definition"] or f"{e['Term']}.",
                          "Purpose": e["purpose"] or "",
                          "Sensitivity": sens, "PII_Category": "",
                          "Critical_Data_Element": "Yes" if e["cde"] else "No",
                          "Abbreviation": _abbrev(e["Term"]),
                          "Suggested_Tags": ";".join(e["tags"]), "Status": "Draft",
                          "Confidence": "High", "Suggested_Reason": "From existing glossary (not in scan)",
                          "LLM_Enriched": "No"})
    return rows + added, {"glossary": g.get("name"), "matched": matched,
                          "added": len(added), "export_terms": len(idx)}

def _kept_rows(rows):
    """Yield only the rows the reviewer marked to keep."""
    return [r for r in rows if str(r.get("Keep", "Y")).lower() in ("y", "yes", "true", "1")]

def _parse_source(src):
    """Resolve a Source_Column into a physical data element (schema/table/column or object)."""
    src = (src or "").strip()
    if not src or src.startswith("glossary:"):
        return None
    if "/" in src and "." not in src.split("/")[0]:
        parts = src.split("/")
        return {"schema_name": parts[0], "table_name": "/".join(parts[1:-1]),
                "column_name": parts[-1], "entity_type": "OBJECT"}
    p = src.split(".")
    if len(p) >= 3:
        return {"schema_name": p[0], "table_name": p[1], "column_name": ".".join(p[2:]),
                "entity_type": "COLUMN"}
    if len(p) == 2:
        return {"schema_name": "", "table_name": p[0], "column_name": p[1], "entity_type": "COLUMN"}
    return None

def _gloss_ns(glossary_name):
    """The UUID5 namespace for one glossary — every id below derives from it, so the
       term/glossary ids are deterministic from names alone (no PDC round-trip)."""
    return uuid.uuid5(uuid.NAMESPACE_DNS, "suggested-glossary:" + glossary_name)


def det_glossary_id(glossary_name):
    """The glossary's id (== every term's rootId == the businessTerm glossaryId)."""
    return str(uuid.uuid5(_gloss_ns(glossary_name), "glossary:" + glossary_name))


def det_term_id(glossary_name, category, term):
    """A term's id, matching its `_id` in the generated glossary JSONL (which PDC
       preserves on import). Category is part of the key, mirroring the JSONL build."""
    return str(uuid.uuid5(_gloss_ns(glossary_name), f"term:{category}/{term}"))


# --- Selective mapping policy -------------------------------------------------
# Not every suggested term should become a PDC data-element association. Linking
# every column pollutes governance, lineage and search, and does nothing for the
# Trust Score (whose glossary-term input is binary - presence, not volume). So a
# term is mapped to its column only when it clears a relevance bar: it is a Critical
# Data Element, it is PII, or it has real evidence behind it (a DDL comment, a key,
# or a profiling hit => High/Medium confidence). Low-confidence, name-templated
# columns are left unmapped. The steward can override per row with a "Map" cell
# (Y/N); the whole gate can be tuned or disabled via the policy.
CONF_RANK = {"low": 0, "medium": 1, "high": 2}

DEFAULT_MAP_POLICY = {
    "mode": "policy",            # "policy" = selective gate (default); "all" = legacy link-everything
    "min_confidence": "medium",  # map terms with at least this evidence; weaker (Low) are skipped
    "always_cde": True,          # always map Critical Data Elements, whatever their confidence
    "always_pii": True,          # always map PII columns, whatever their confidence
}

def should_map_link(row, policy=None):
    """Decide whether a reviewed term should be linked to its data element.
       Returns (map?, reason). A per-row "Map" cell (Y/N) always wins; otherwise the
       policy gates on CDE / PII / confidence. 'No match' is a valid, deliberate
       outcome - see DEFAULT_MAP_POLICY."""
    pol = {**DEFAULT_MAP_POLICY, **(policy or {})}
    ov = str(row.get("Map", "")).strip().lower()
    if ov in ("n", "no", "false", "0", "skip"):
        return False, "steward set Map=No"
    if ov in ("y", "yes", "true", "1", "map"):
        return True, "steward set Map=Yes"
    if pol.get("mode") == "all":
        return True, "policy: map all"
    cde = str(row.get("Critical_Data_Element", "No")).strip().lower() == "yes"
    pii = bool(str(row.get("PII_Category", "")).strip())
    if pol.get("always_cde") and cde:
        return True, "Critical Data Element"
    if pol.get("always_pii") and pii:
        return True, "PII column"
    conf = str(row.get("Confidence", "Low")).strip().lower()
    floor = str(pol.get("min_confidence", "medium")).strip().lower()
    if CONF_RANK.get(conf, 0) >= CONF_RANK.get(floor, 1):
        return True, f"{row.get('Confidence', 'Low')} confidence"
    return False, f"{row.get('Confidence', 'Low')} confidence, not CDE/PII"

def _row_real_sources(row):
    """The physical (non-glossary) source columns a row would actually link to."""
    return [s.strip() for s in str(row.get("Source_Column", "")).split(";")
            if _parse_source(s.strip())]

def map_breakdown(rows, policy=None):
    """Explain the gate: which kept terms get mapped to columns and which are held
       back (and why). Drives the steward-facing summary so selectivity is visible,
       not silent."""
    mapped, skipped = [], []
    for r in _kept_rows(rows):
        srcs = _row_real_sources(r)
        ok, why = should_map_link(r, policy)
        item = {"term": r.get("Term", ""), "category": r.get("Category", ""),
                "confidence": r.get("Confidence", ""),
                "cde": r.get("Critical_Data_Element", "No"),
                "pii": r.get("PII_Category", ""), "columns": len(srcs), "reason": why}
        if not srcs:
            item["reason"] = "conceptual / glossary-only term (no physical column)"
            skipped.append(item)
        elif ok:
            mapped.append(item)
        else:
            skipped.append(item)
    return {"mapped": mapped, "skipped": skipped,
            "mapped_count": len(mapped), "skipped_count": len(skipped)}

def data_element_links(rows, glossary_name="Business Glossary", quality_weights=None, with_quality=True, policy=None):
    """Map each kept term to the physical column(s) it came from — the Data Element
       associations (term <-> column) keyed by schema/table/column for bulk assignment.
       Each link carries the column's own scan-suggested rating and DQ qualityScore
       (the latter recomputed here so weights can be tuned without re-scanning).

       Selectivity: only rows the policy keeps are linked (see should_map_link), so
       low-value, non-CDE, non-PII columns are not auto-associated. Pass
       policy={"mode": "all"} to restore the legacy link-every-term behaviour."""
    links = []
    for r in _kept_rows(rows):
        keep, _why = should_map_link(r, policy)
        if not keep:
            continue
        ratings_map = r.get("Source_Ratings") or {}
        fallback = int(r.get("Suggested_Rating", 0) or 0)
        qdims_map = r.get("Source_Quality_Dims") or {}
        keys_map = r.get("Source_Keys") or {}
        for sc in str(r.get("Source_Column", "")).split(";"):
            sc_key = sc.strip()
            de = _parse_source(sc_key)
            if not de:
                continue
            # each physical column carries its own scan-suggested rating; fall back
            # to the term's representative rating if the per-column value is missing
            rating = int(ratings_map.get(sc_key, fallback) or fallback or 0)
            # DQ score from this column's own dimensions, under the chosen weights
            quality = None
            qd = qdims_map.get(sc_key)
            if with_quality and qd:
                quality = quality_score_column(completeness=qd.get("c"), uniqueness=qd.get("u"),
                                               validity=qd.get("v"), expect_unique=qd.get("eu"),
                                               notnull=qd.get("nn"), weights=quality_weights)
            links.append({**de, "business_term": r["Term"], "glossary": glossary_name,
                          "category": r.get("Category", ""), "sensitivity": r.get("Sensitivity", ""),
                          "critical_data_element": r.get("Critical_Data_Element", "No"),
                          "rating": rating, "quality": quality,
                          "definition": (r.get("Definition") or "").strip(),
                          "keys": keys_map.get(sc_key)})
    return links

DE_COLS = ["schema_name", "table_name", "column_name", "entity_type", "business_term",
           "glossary", "category", "sensitivity", "critical_data_element"]

def links_to_csv(links):
    """Render Data-Element links as bulk-assignment CSV."""
    import csv, io
    buf = io.StringIO()
    w = csv.DictWriter(buf, fieldnames=DE_COLS, extrasaction="ignore")
    w.writeheader()
    for l in links:
        w.writerow(l)
    return buf.getvalue()

def links_to_api_json(links, glossary_name="Business Glossary", lineage_verified=True, rating=0):
    """Trust-Score-ready association objects, one per column, shaped for PDC's
       data-collections API: businessTerms + features(isLineageVerified, rating,
       sensitivity, isCriticalDataElement). A linked term + verified lineage + a
       rating + a quality metric are the inputs PDC's Trust Score draws on.

       Each column carries its own scan-suggested rating (link['rating']); when
       multiple terms map to one column the highest suggestion wins. Pass a
       non-zero `rating` to override every column with one fixed value instead."""
    by_col = {}
    for l in links:
        key = (l["schema_name"], l["table_name"], l["column_name"], l["entity_type"])
        rec = by_col.setdefault(key, {
            "type": l["entity_type"], "schemaName": l["schema_name"],
            "tableName": l["table_name"], "columnName": l["column_name"],
            "attributes": {"businessTerms": [],
                           "features": {"sensitivity": l.get("sensitivity", ""),
                                        "isCriticalDataElement": str(l.get("critical_data_element", "No")).lower() == "yes",
                                        "isLineageVerified": bool(lineage_verified)}}})
        # rating: explicit global override, else the highest scan suggestion for the column
        col_rating = int(rating) if rating else int(l.get("rating") or 0)
        if col_rating:
            cur = (rec["attributes"]["features"].get("rating") or {}).get("value", 0)
            if col_rating >= cur:
                rec["attributes"]["features"]["rating"] = {"value": col_rating}
        # qualityScore: the Data Quality input (0-100). Highest scan suggestion wins
        # when several terms map to one column. PDC records an externally-set value
        # as a MANUAL quality metric (which is what we want now PDQ is retired).
        q = l.get("quality")
        if q is not None:
            curq = rec["attributes"]["features"].get("qualityScore")
            if curq is None or int(q) >= int(curq):
                rec["attributes"]["features"]["qualityScore"] = int(q)
        # Stamp the term's id and the glossaryId deterministically — they are the
        # SAME UUID5s written into the glossary JSONL (which PDC preserves on import),
        # so the link is born fully glossary-bound (id + glossaryId) with no PDC
        # round-trip. Resolve then only has to confirm, and Apply writes a real link
        # instead of attaching by name (which leaves the Glossary column as "—").
        # entity description: the steward's reviewed definition (PATCHable via
        # attributes.info.description; Apply decides fill-vs-overwrite)
        if l.get("definition") and "info" not in rec["attributes"]:
            rec["attributes"]["info"] = {"description": l["definition"]}
        # PK/FK facts -> attributes.extended. The built-in Is Primary/Foreign Key
        # property (metadata.column.*) is harvest-owned and rejected by the public
        # PATCH schema; extended is the API's writable free-form block, so the
        # scan's own key detection is recorded there.
        kk = l.get("keys")
        if isinstance(kk, dict) and (kk.get("pk") or kk.get("fk")):
            ext = rec["attributes"].setdefault("extended", {})
            ext["isPrimaryKey"] = bool(kk.get("pk"))
            ext["isForeignKey"] = bool(kk.get("fk"))
            if kk.get("ref"):
                ext["references"] = kk["ref"]
        gname = l.get("glossary", glossary_name) or glossary_name
        rec["attributes"]["businessTerms"].append(
            {"name": l["business_term"], "glossary": gname,
             "id": det_term_id(gname, l.get("category", ""), l["business_term"]),
             "glossaryId": det_glossary_id(gname)})
    return list(by_col.values())

def table_term_directory(rows, glossary_name="Business Glossary"):
    """{table_name(lower): term info} for the table-level record terms, so Apply
       can bind each table's OWN businessTerm — plus its description and
       sensitivity — onto the TABLE entity. That is the Trust Score's
       "glossary term assigned" input at table level, automated (it was a
       documented manual steward step before 1.8.6). Ids are the same
       deterministic UUID5s the glossary JSONL carries, so the link is born
       glossary-bound once the glossary is imported."""
    out = {}
    for r in _kept_rows(rows):
        t = (r.get("Source_Table") or "").strip()
        term = (r.get("Term") or "").strip()
        if not t or not term:
            continue
        out[t.lower()] = {
            "name": term,
            "id": det_term_id(glossary_name, r.get("Category", ""), term),
            "glossaryId": det_glossary_id(glossary_name),
            "description": (r.get("Definition") or "").strip(),
            "sensitivity": (r.get("Sensitivity") or "").strip().upper(),
        }
    return out

def glossary_to_rows(jsonl_text):
    """Load an exported glossary directly as editable review rows (round-trip / review)."""
    g = parse_glossary(jsonl_text)
    rows = []
    for e in g["terms"].values():
        sens = e["sensitivity"] or "LOW"
        rows.append({"Keep": "Y", "Category": e["Category"], "Term": e["Term"],
                     "Source_Column": f"glossary:{g.get('name') or 'imported'}",
                     "Definition": e["definition"] or f"{e['Term']}.",
                     "Purpose": e["purpose"] or "", "Sensitivity": sens, "PII_Category": "",
                     "Critical_Data_Element": "Yes" if e["cde"] else "No",
                     "Abbreviation": _abbrev(e["Term"]), "Suggested_Tags": ";".join(e["tags"]),
                     "Status": "Draft", "Confidence": "High",
                     "Suggested_Reason": "Loaded from glossary export", "LLM_Enriched": "No"})
    return rows, {"glossary": g.get("name"), "terms": len(rows),
                  "categories": len({r["Category"] for r in rows})}

# ----------------------------------------------------------------- GENERATE
def _lex(text):
    """Wrap plain text as a minimal Lexical-JSON string (PDC's rich-text format)."""
    if not text:
        return None
    obj = {"root": {"children": [{"children": [{"detail": 0, "format": 0, "mode": "normal",
            "style": "", "text": str(text), "type": "text", "version": 1}],
            "direction": "ltr", "format": "", "indent": 0, "type": "paragraph", "version": 1}],
            "direction": "ltr", "format": "", "indent": 0, "type": "root", "version": 1}}
    return json.dumps(obj, ensure_ascii=False)

SENS_RANK = {"LOW": 0, "MEDIUM": 1, "HIGH": 2}
RANK_SENS = {0: "LOW", 1: "MEDIUM", 2: "HIGH"}
REVIEW_TIME = "T00:00:00.000Z"   # PDC stores reviewedAt at midnight UTC

def _classification(sens):
    """Map a sensitivity level to PDC's classification label (HIGH->Company Confidential, MEDIUM->Private, else Public)."""
    return ("Company Confidential" if sens == "HIGH"
            else "Private" if sens == "MEDIUM" else "Public")

def _kept(r):
    """True when a row is marked kept (Keep in y/yes/true/1)."""
    return str(r.get("Keep", "Y")).strip().lower() in ("y", "yes", "true", "1")

def _people_block(people, rating, review_iso, created_override=""):
    """Build the reusable injection pieces for one people-scope (glossary-wide or
       per-category). `people` = {owner, custodian, businessSteward, stakeholders}.
       Safe/empty when nothing is set (output then matches the no-governance case)."""
    people = people or {}
    owner = (people.get("owner") or "").strip()
    custodian = (people.get("custodian") or "").strip()
    steward = (people.get("businessSteward") or "").strip()

    info_people = {}
    if owner:     info_people["owner"] = owner
    if custodian: info_people["custodian"] = custodian
    if steward:   info_people["businessSteward"] = steward

    features_extra = {}
    rater = steward or owner or custodian
    if rating and rater:
        features_extra["rating"] = {"value": rating, "users": {rater: rating}}

    attr_extra = {}
    clean_sh = []
    for s in (people.get("stakeholders") or []):
        sid = (s.get("id") or "").strip()
        if not sid:
            continue
        clean_sh.append({"roles": s.get("roles") or ["Steward"],
                         "name": s.get("name") or "", "id": sid, "email": s.get("email") or ""})
    if clean_sh:
        attr_extra["stakeholders"] = clean_sh
    if review_iso:
        attr_extra["reviewedAt"] = review_iso

    created_by = (created_override or steward or owner or "suggester")
    updated_by = (steward or owner or "suggester")
    return info_people, features_extra, attr_extra, created_by, updated_by


def _merge_people(base, over):
    """Per-category override: replace owner/custodian/businessSteward/stakeholders
       individually when the override supplies them, else inherit from base."""
    out = dict(base or {})
    over = over or {}
    for k in ("owner", "custodian", "businessSteward"):
        if (over.get(k) or "").strip():
            out[k] = over[k]
    if over.get("stakeholders"):
        out["stakeholders"] = over["stakeholders"]
    return out


def _cat_effective(over, g_status, g_rating, g_review_iso):
    """Resolve a category's effective status / rating / reviewed-date. Each field
       falls back to the glossary-wide value unless the category overrides it.
       An empty string means 'use default'; rating '0' is a real override (None)."""
    over = over or {}
    status = (str(over.get("status") or "")).strip() or g_status
    rraw = over.get("rating", "")
    if str(rraw).strip() != "":
        try:
            rating = int(rraw)
        except (TypeError, ValueError):
            rating = g_rating
    else:
        rating = g_rating
    rv = (over.get("reviewedAt") or "").strip()
    review_iso = ((rv if "T" in rv else rv + REVIEW_TIME) if rv else g_review_iso)
    return status, rating, review_iso


def to_jsonl_records(rows, glossary_name="Business Glossary (Suggested)", governance=None):
    """Build PDC glossary import records (glossary/category/term objects) from review rows."""
    gov = governance or {}
    status = (gov.get("status") or "Draft").strip() or "Draft"
    domain = (gov.get("domain") or DOMAIN).strip() or DOMAIN  # glossary-wide PDC classifier
    apply_cat = gov.get("applyToCategories", True)
    created_override = (gov.get("createdBy") or "").strip()
    try:
        rating = int(gov.get("rating") or 0)
    except (TypeError, ValueError):
        rating = 0
    rv = (gov.get("reviewedAt") or "").strip()
    review_iso = (rv if "T" in rv else (rv + REVIEW_TIME)) if rv else ""

    # default people scope: explicit "default", else legacy top-level fields
    default_people = gov.get("default") or {
        "owner": gov.get("owner", ""), "custodian": gov.get("custodian", ""),
        "businessSteward": gov.get("businessSteward", ""), "stakeholders": gov.get("stakeholders", [])}
    cat_overrides = gov.get("categories") or {}

    ns = uuid.uuid5(uuid.NAMESPACE_DNS, "suggested-glossary:" + glossary_name)
    root = det_glossary_id(glossary_name)

    g_info, g_feat, g_attr, g_cby, g_uby = _people_block(default_people, rating, review_iso, created_override)
    gloss_info = {"status": status}
    gloss_info.update(g_info)
    recs = [{"createdAt": GEN_TS, "fqdn": glossary_name, "rootId": root, "createdBy": g_cby,
             "name": glossary_name, "attributes": {"isSoftCreated": False, "info": gloss_info},
             "type": "glossary", "updatedAt": GEN_TS, "resourceId": "null", "_id": root, "sort": None}]

    cats, cat_id, cat_rows = [], {}, {}
    for r in rows:
        if not _kept(r):
            continue
        cat_rows.setdefault(r["Category"], []).append(r)
        if r["Category"] not in cats:
            cats.append(r["Category"])

    # resolve each category's effective people-block + status once; terms inherit theirs
    cat_block = {}
    cat_status = {}
    for cat in cats:
        ov = cat_overrides.get(cat)
        c_status, c_rating, c_review = _cat_effective(ov, status, rating, review_iso)
        cat_status[cat] = c_status
        cat_people = _merge_people(default_people, ov)
        cat_block[cat] = _people_block(cat_people, c_rating, c_review, created_override)

    for cat in cats:
        cid = str(uuid.uuid5(ns, "category:" + cat)); cat_id[cat] = cid
        infoP, featX, attrX, cby, uby = cat_block[cat]
        csens = RANK_SENS[max((SENS_RANK.get(x["Sensitivity"], 0) for x in cat_rows[cat]), default=0)]
        cinfo = {"domain": domain,
                 "definition": _lex(f"{cat} terms in the {glossary_name} business glossary."),
                 "classification": _classification(csens), "status": cat_status[cat],
                 "purpose": _lex(f"Groups {cat.lower()} business terms for governance and discovery.")}
        cattrs = {"features": {"sensitivity": csens}, "isSoftCreated": False, "info": cinfo}
        if apply_cat:
            cinfo.update(infoP)
            cattrs["features"].update(featX)
            cattrs.update(dict(attrX))
        else:
            cattrs = {"isSoftCreated": False, "info": {"domain": domain, "status": cat_status[cat]}}
        recs.append({"createdAt": GEN_TS, "updatedBy": uby, "fqdn": f"{glossary_name}/{cat}",
                     "rootId": root, "createdBy": cby, "name": cat,
                     "attributes": cattrs, "type": "category", "parentId": root,
                     "updatedAt": GEN_TS, "resourceId": "null", "_id": cid, "sort": None})

    for r in rows:
        if not _kept(r):
            continue
        cat = r["Category"]
        infoP, featX, attrX, cby, uby = cat_block[cat]
        tid = det_term_id(glossary_name, cat, r['Term'])
        sens = r["Sensitivity"]
        info = {"domain": domain, "definition": _lex(r["Definition"]),
                "classification": _classification(sens),
                "status": cat_status[cat], "purpose": _lex(r.get("Purpose") or f"Suggested from {r['Source_Column']}.")}
        info.update(infoP)
        if r.get("Abbreviation"):
            info["abbreviation"] = r["Abbreviation"]
        features = {"sensitivity": sens,
                    "isCriticalDataElement": str(r["Critical_Data_Element"]).lower() == "yes"}
        features.update(featX)
        attrs = {"features": features, "isSoftCreated": False, "info": info}
        attrs.update(dict(attrX))
        if r.get("Suggested_Tags"):
            attrs["tags"] = [{"name": t.strip().lower()}
                             for t in r["Suggested_Tags"].split(";") if t.strip()]
        recs.append({"createdAt": GEN_TS, "updatedBy": uby,
                     "fqdn": f"{glossary_name}/{cat}/{r['Term']}", "rootId": root,
                     "createdBy": cby, "name": r["Term"], "attributes": attrs,
                     "type": "term", "parentId": cat_id[cat], "updatedAt": GEN_TS,
                     "resourceId": "null", "_id": tid, "sort": None})
    return recs

def records_to_jsonl(recs):
    """Serialise glossary records to JSONL (one JSON object per line)."""
    return "\n".join(json.dumps(r, ensure_ascii=False) for r in recs) + "\n"


# --------------------------------------------------------------------------- #
#  Action "verdict" checks — each returns {title, rows[], issues[], tone, verdict}
#  so the UI can show "what came back + a plain-English verdict" (like the PDC
#  Resolve probe) for Generate JSONL, Scan, and Document Discovery.
# --------------------------------------------------------------------------- #
def _trunc(items, n=6):
    items = list(items)
    return ", ".join(items[:n]) + ("…" if len(items) > n else "")


def glossary_build_check(rows, recs, glossary_name):
    """Build-time sanity check for the glossary JSONL: counts, plus the issues that
       actually bite on import or Resolve — id collisions (same term twice in a
       category share one UUID5 id), names that repeat across categories (ambiguous
       for name-based Resolve), and missing category/definition."""
    from collections import Counter
    kept = [r for r in rows if _kept(r)]
    terms = [r for r in recs if r.get("type") == "term"]
    cats = [r for r in recs if r.get("type") == "category"]

    no_def = [r.get("Term", "") for r in kept if not str(r.get("Definition", "")).strip()]
    no_cat = [r.get("Term", "") for r in kept if not str(r.get("Category", "")).strip()]
    pair_ct = Counter((r.get("Category", ""), r.get("Term", "")) for r in kept)
    dup_pairs = sorted(f"{(c or '—')} / {t}" for (c, t), n in pair_ct.items() if n > 1)
    name_ct = Counter(r.get("Term", "") for r in kept)
    dup_names = sorted(t for t, n in name_ct.items() if n > 1)

    rows_out = [
        {"label": "Glossary", "value": glossary_name},
        {"label": "Lines", "value": f"{len(recs)} ({len(cats)} categories, {len(terms)} terms)"},
        {"label": "Kept / dropped", "value": f"{len(kept)} / {len(rows) - len(kept)}"},
    ]
    issues = []
    if dup_pairs:
        issues.append({"tone": "bad", "text": f"{len(dup_pairs)} term(s) duplicated within a category — "
                       f"these share one generated id and collide on import (one overwrites the other): {_trunc(dup_pairs)}"})
    if dup_names:
        issues.append({"tone": "warn", "text": f"{len(dup_names)} term name(s) repeat across categories — "
                       f"name-based Resolve can't tell them apart, so a column may link to the wrong one: {_trunc(dup_names)}"})
    if no_cat:
        issues.append({"tone": "warn", "text": f"{len(no_cat)} term(s) have no category — they import under 'Unassigned': {_trunc(no_cat)}"})
    if no_def:
        issues.append({"tone": "warn", "text": f"{len(no_def)} term(s) have no definition: {_trunc(no_def)}"})

    tone = "bad" if dup_pairs else ("warn" if issues else "ok")
    verdict = ({"ok": f"All {len(terms)} terms are clean — import this JSONL in PDC (Glossary → Actions → Import), then Resolve & Apply.",
                "warn": "Importable, but the notes above can cause ambiguous links or Unassigned terms. Fix them in the table for clean links.",
                "bad": "Duplicate terms in the same category lose data on import — rename or remove the duplicates before importing."})[tone]
    return {"title": "Build check", "rows": rows_out, "issues": issues, "tone": tone, "verdict": verdict}


def scan_check(rows, scanned, pk_cols=0, fk_cols=0):
    """Verdict for a scan: what the catalog saw and what's worth a look before
       generating (no PKs, no DQ, lots of low-confidence templated terms)."""
    from collections import Counter
    sens = Counter(r.get("Sensitivity", "") for r in rows)
    conf = Counter(r.get("Confidence", "") for r in rows)
    cde = sum(1 for r in rows if str(r.get("Critical_Data_Element", "")).lower() == "yes")
    pii = sum(1 for r in rows if str(r.get("PII_Category", "")).strip())
    dq = sum(1 for r in rows if r.get("Suggested_Quality") is not None)
    is_db = scanned.get("tables") is not None

    rows_out = []
    if is_db:
        rows_out.append({"label": "Scanned", "value": f"{scanned.get('tables', 0)} tables · {scanned.get('columns', 0)} columns"})
    if scanned.get("objects") is not None:
        rows_out.append({"label": "Scanned", "value": f"{scanned.get('objects', 0)} files · {scanned.get('folders', 0)} folders"})
    rows_out.append({"label": "Terms suggested", "value": str(len(rows))})
    rows_out.append({"label": "Sensitivity", "value": f"HIGH {sens.get('HIGH', 0)} · MED {sens.get('MEDIUM', 0)} · LOW {sens.get('LOW', 0)}"})
    rows_out.append({"label": "CDE / PII", "value": f"{cde} critical · {pii} PII"})
    if is_db:
        rows_out.append({"label": "Keys detected", "value": f"{pk_cols} PK · {fk_cols} FK"})
    rows_out.append({"label": "Confidence", "value": f"High {conf.get('High', 0)} · Med {conf.get('Medium', 0)} · Low {conf.get('Low', 0)}"})
    if is_db:
        rows_out.append({"label": "DQ computed", "value": f"{dq} of {len(rows)} columns"})

    issues = []
    if is_db and not pk_cols:
        issues.append({"tone": "warn", "text": "No primary keys detected — PDC's 'Is Primary Key' comes from the DB catalog scan, "
                       "so re-catalog the source (and check the JDBC driver) if you expect PKs."})
    if is_db and dq == 0:
        issues.append({"tone": "warn", "text": "No Data Quality computed — turn on profiling for this connection so each column gets a "
                       "DQ score (one of the four Trust-Score inputs)."})
    low = conf.get("Low", 0)
    if low and low > max(1, len(rows) // 2):
        issues.append({"tone": "warn", "text": f"{low} of {len(rows)} terms are Low confidence (templated from the column name) — "
                       "review their definitions before generating."})
    tone = "warn" if issues else "ok"
    verdict = ("Scan looks good — review the suggested terms, then Generate JSONL." if tone == "ok"
               else "Scan complete. The notes above are worth a look before you generate the glossary.")
    return {"title": "Scan check", "rows": rows_out, "issues": issues, "tone": tone, "verdict": verdict}
